from __future__ import annotations

import re
import sys
import time
from html import escape
from io import BytesIO
from pathlib import Path

import pandas as pd
import streamlit as st

LOCAL_PACKAGES = Path(__file__).resolve().parent / ".packages"
if LOCAL_PACKAGES.exists():
    sys.path.insert(0, str(LOCAL_PACKAGES))

from careercompass.agents import (
    evaluate_interview_answer,
    generate_interview_questions,
    run_career_analysis,
)
from careercompass.demo_data import SAMPLE_COURSEWORK, SAMPLE_RESUME
from careercompass.fallbacks import assess_skill_evidence
from careercompass.rag import (
    available_coursework_options,
    available_locations,
    available_target_roles,
)


ACCENT_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap');

:root {
    --cc-navy-950: #07111f;
    --cc-navy-900: #0b1628;
    --cc-navy-800: #12243a;
    --cc-teal-500: #14b8a6;
    --cc-teal-600: #0f766e;
    --cc-blue-500: #2563eb;
    --cc-amber-500: #f59e0b;
    --cc-surface: rgba(255, 255, 255, 0.92);
    --cc-surface-strong: #ffffff;
    --cc-border: rgba(15, 23, 42, 0.1);
    --cc-text: #101827;
    --cc-muted: #64748b;
}

.stApp {
    background:
        linear-gradient(115deg, rgba(20, 184, 166, 0.08) 0 1px, transparent 1px 96px),
        linear-gradient(180deg, #f6fbfd 0%, #eef6f8 52%, #f8fbff 100%);
    color: var(--cc-text);
    font-family: Inter, sans-serif;
}

[data-testid="stSidebar"] {
    background:
        linear-gradient(180deg, var(--cc-navy-950) 0%, var(--cc-navy-900) 58%, #0b3f4a 100%);
    border-right: 1px solid rgba(148, 163, 184, 0.18);
}

[data-testid="stSidebar"] * {
    color: #f8fafc;
}

h1 {
    font-weight: 800;
    letter-spacing: 0;
    color: #101827;
}

h2, h3 {
    letter-spacing: 0;
    color: #1f2937;
}

p, .stCaptionContainer {
    color: var(--cc-muted);
}

.cc-hero {
    position: relative;
    overflow: hidden;
    border: 1px solid rgba(20, 184, 166, 0.2);
    border-radius: 16px;
    padding: 28px 30px 30px 30px;
    margin: 8px 0 22px 0;
    background:
        radial-gradient(circle at 86% 32%, rgba(20, 184, 166, 0.34), transparent 28%),
        radial-gradient(circle at 70% 8%, rgba(37, 99, 235, 0.16), transparent 24%),
        linear-gradient(135deg, rgba(7, 17, 31, 0.98), rgba(13, 37, 58, 0.97) 58%, rgba(15, 76, 84, 0.94));
    box-shadow: 0 22px 55px rgba(15, 23, 42, 0.18);
}

.cc-hero::before {
    content: "";
    position: absolute;
    inset: 0;
    background:
        linear-gradient(90deg, rgba(148, 163, 184, 0.08) 1px, transparent 1px),
        linear-gradient(0deg, rgba(148, 163, 184, 0.06) 1px, transparent 1px),
        radial-gradient(circle at 76% 58%, transparent 0 18px, rgba(20, 184, 166, 0.34) 19px, transparent 21px);
    background-size: 52px 52px, 52px 52px, auto;
    opacity: 0.72;
    pointer-events: none;
}

.cc-hero::after {
    content: "";
    position: absolute;
    inset: 24px 22px auto auto;
    width: 320px;
    height: 150px;
    border-top: 1px dashed rgba(125, 211, 252, 0.36);
    border-right: 1px dashed rgba(20, 184, 166, 0.34);
    border-radius: 0 28px 0 0;
    transform: skewX(-14deg);
    pointer-events: none;
}

.cc-hero-grid {
    position: relative;
    z-index: 2;
    display: grid;
    grid-template-columns: 1fr auto;
    gap: 24px;
    align-items: end;
}

.cc-hero-kicker {
    color: #99f6e4;
    font-size: 0.78rem;
    font-weight: 800;
    letter-spacing: 0.08em;
    margin-bottom: 8px;
    text-transform: uppercase;
}

.cc-hero h1 {
    margin: 0 0 8px 0;
    color: #f8fafc;
    font-size: 3.4rem;
    line-height: 0.95;
}

.cc-hero p {
    max-width: 760px;
    margin: 0;
    color: #cbd5e1;
    font-size: 1.02rem;
    line-height: 1.55;
}

.cc-compass-mark {
    position: relative;
    width: 92px;
    height: 92px;
    border-radius: 50%;
    border: 1px solid rgba(20, 184, 166, 0.46);
    display: grid;
    place-items: center;
    background: rgba(15, 23, 42, 0.28);
    box-shadow:
        inset 0 0 0 12px rgba(20, 184, 166, 0.08),
        inset 0 0 0 32px rgba(255, 255, 255, 0.02),
        0 18px 34px rgba(0, 0, 0, 0.22);
}

.cc-compass-mark::before {
    content: "";
    position: absolute;
    inset: 18px;
    border: 1px solid rgba(153, 246, 228, 0.42);
    border-radius: 50%;
    background:
        linear-gradient(90deg, transparent 49%, rgba(226, 232, 240, 0.45) 49% 51%, transparent 51%),
        linear-gradient(0deg, transparent 49%, rgba(226, 232, 240, 0.45) 49% 51%, transparent 51%);
}

.cc-compass-mark::after {
    content: "";
    position: absolute;
    left: 50%;
    top: 50%;
    width: 16px;
    height: 58px;
    background: linear-gradient(180deg, #f59e0b 0 46%, #e2e8f0 48% 100%);
    clip-path: polygon(50% 0, 100% 46%, 59% 50%, 50% 100%, 41% 50%, 0 46%);
    transform: translate(-50%, -50%) rotate(38deg);
    filter: drop-shadow(0 5px 10px rgba(0, 0, 0, 0.28));
}

.cc-hero-metrics {
    position: relative;
    z-index: 2;
    display: grid;
    grid-template-columns: repeat(4, minmax(0, 1fr));
    gap: 12px;
    margin-top: 22px;
}

.cc-hero-metric {
    min-height: 112px;
    border: 1px solid rgba(226, 232, 240, 0.72);
    border-radius: 10px;
    padding: 14px 16px;
    background: rgba(255, 255, 255, 0.96);
    box-shadow: 0 16px 34px rgba(2, 6, 23, 0.14);
}

.cc-hero-metric span {
    display: block;
    color: #475569;
    font-size: 0.75rem;
    font-weight: 800;
    margin-bottom: 8px;
    text-transform: uppercase;
}

.cc-hero-metric strong {
    display: block;
    color: #0b1628;
    font-size: 1.72rem;
    line-height: 1.1;
}

.cc-hero-metric em {
    display: block;
    color: #0f766e;
    font-size: 0.78rem;
    font-style: normal;
    font-weight: 800;
    margin-top: 7px;
}

.cc-hero-metric.is-alert em {
    color: #b45309;
}

.cc-chip-row {
    position: relative;
    z-index: 2;
    display: flex;
    flex-wrap: wrap;
    gap: 10px;
    margin-top: 18px;
}

.cc-chip {
    display: inline-flex;
    align-items: center;
    gap: 8px;
    padding: 8px 11px;
    border: 1px solid rgba(148, 163, 184, 0.24);
    border-radius: 999px;
    background: rgba(15, 23, 42, 0.48);
    color: #e2e8f0;
    font-size: 0.8rem;
    font-weight: 700;
}

.cc-chip::before {
    content: "";
    width: 7px;
    height: 7px;
    border-radius: 50%;
    background: var(--cc-teal-500);
    box-shadow: 0 0 0 3px rgba(20, 184, 166, 0.16);
}

.cc-sidebar-brand {
    display: flex;
    gap: 12px;
    align-items: center;
    padding: 10px 0 8px 0;
}

.cc-sidebar-compass {
    position: relative;
    flex: 0 0 38px;
    width: 38px;
    height: 38px;
    border-radius: 50%;
    border: 1px solid rgba(153, 246, 228, 0.52);
    box-shadow: inset 0 0 0 7px rgba(20, 184, 166, 0.08);
}

.cc-sidebar-compass::before {
    content: "";
    position: absolute;
    inset: 8px;
    border: 1px solid rgba(203, 213, 225, 0.44);
    border-radius: 50%;
    background:
        linear-gradient(90deg, transparent 48%, rgba(226, 232, 240, 0.55) 48% 52%, transparent 52%),
        linear-gradient(0deg, transparent 48%, rgba(226, 232, 240, 0.55) 48% 52%, transparent 52%);
}

.cc-sidebar-compass::after {
    content: "";
    position: absolute;
    left: 50%;
    top: 50%;
    width: 7px;
    height: 24px;
    background: linear-gradient(180deg, #f59e0b 0 46%, #e2e8f0 48% 100%);
    clip-path: polygon(50% 0, 100% 46%, 58% 50%, 50% 100%, 42% 50%, 0 46%);
    transform: translate(-50%, -50%) rotate(38deg);
}

.cc-sidebar-brand div:last-child {
    min-width: 0;
}

.cc-sidebar-brand strong {
    display: block;
    font-size: 1.28rem;
    letter-spacing: 0;
}

.cc-sidebar-brand span {
    display: block;
    margin-top: 4px;
    color: #94a3b8 !important;
    font-size: 0.82rem;
}

.cc-sidebar-label {
    margin: 22px 0 12px 0;
    color: #67e8f9 !important;
    font-size: 0.74rem;
    font-weight: 800;
    letter-spacing: 0.08em;
    text-transform: uppercase;
}

.cc-sidebar-route {
    margin-top: 0;
    padding-left: 12px;
    border-left: 1px solid rgba(20, 184, 166, 0.32);
}

.cc-sidebar-step {
    margin: 0 0 16px 0;
}

.cc-sidebar-step strong {
    display: block;
    font-size: 0.85rem;
}

.cc-sidebar-step span {
    display: block;
    color: #94a3b8 !important;
    font-size: 0.75rem;
    margin-top: 2px;
}

.cc-page-signal {
    border-left: 4px solid var(--cc-teal-500);
    padding: 10px 14px;
    margin: 4px 0 18px 0;
    border-radius: 8px;
    background: rgba(240, 253, 250, 0.78);
    color: #0f766e;
    font-weight: 700;
    font-size: 0.9rem;
}

.cc-kicker {
    color: #0f766e;
    font-size: 0.82rem;
    font-weight: 800;
    letter-spacing: 0.08em;
    text-transform: uppercase;
    margin-bottom: 6px;
}

.cc-card {
    border: 1px solid var(--cc-border);
    border-radius: 8px;
    padding: 18px;
    background: var(--cc-surface);
    box-shadow: 0 12px 30px rgba(15, 23, 42, 0.06);
}

.cc-status-grid {
    display: grid;
    grid-template-columns: repeat(5, minmax(0, 1fr));
    gap: 10px;
    margin: 10px 0 18px 0;
}

.cc-status {
    border: 1px solid rgba(15, 23, 42, 0.08);
    border-radius: 8px;
    padding: 12px;
    background: rgba(255, 255, 255, 0.78);
    min-height: 112px;
}

.cc-status strong {
    display: block;
    font-size: 0.9rem;
    margin-bottom: 4px;
}

.cc-status span {
    display: inline-block;
    color: #0f766e;
    font-size: 0.74rem;
    font-weight: 800;
    margin-bottom: 8px;
    text-transform: uppercase;
}

.cc-status p {
    font-size: 0.82rem;
    line-height: 1.35;
    margin: 0;
}

.cc-demo-cue {
    border-left: 5px solid #0f766e;
    border-radius: 8px;
    padding: 14px 16px;
    background: rgba(240, 253, 250, 0.9);
    margin: 8px 0 18px 0;
}

.cc-demo-cue strong {
    color: #0f766e;
}

@media (max-width: 1100px) {
    .cc-status-grid {
        grid-template-columns: repeat(2, minmax(0, 1fr));
    }
    .cc-hero-metrics {
        grid-template-columns: repeat(2, minmax(0, 1fr));
    }
}

div[data-testid="stMetric"] {
    border: 1px solid rgba(20, 184, 166, 0.16);
    border-radius: 10px;
    padding: 16px 18px;
    background: linear-gradient(180deg, rgba(255, 255, 255, 0.96), rgba(248, 250, 252, 0.9));
    box-shadow: 0 12px 32px rgba(15, 23, 42, 0.07);
}

div[data-testid="stMetric"] [data-testid="stMetricLabel"] {
    color: var(--cc-muted);
    font-weight: 700;
}

div[data-testid="stMetric"] [data-testid="stMetricValue"] {
    color: var(--cc-navy-900);
    font-weight: 800;
}

.stButton > button,
.stDownloadButton > button {
    border-radius: 10px;
    border: 1px solid rgba(15, 118, 110, 0.18);
    color: #0f172a !important;
    font-weight: 700;
    min-height: 2.65rem;
    box-shadow: 0 8px 20px rgba(15, 23, 42, 0.06);
}

.stButton > button p,
.stDownloadButton > button p {
    color: inherit !important;
}

.stButton > button[kind="primary"] {
    background: linear-gradient(90deg, var(--cc-teal-600), var(--cc-blue-500));
    border: 0;
    color: #ffffff !important;
    text-shadow: 0 1px 1px rgba(15, 23, 42, 0.24);
}

.stButton > button[kind="primary"] *,
.stButton > button[kind="primary"] p {
    color: #ffffff !important;
}

.stButton > button[kind="primary"]:hover {
    background: linear-gradient(90deg, #0d9488, #1d4ed8);
    border: 0;
    color: #ffffff !important;
}

.stButton > button[kind="primary"]:disabled,
.stButton > button[kind="primary"]:disabled * {
    color: rgba(255, 255, 255, 0.72) !important;
    opacity: 1;
}

.stRadio div[role="radiogroup"] {
    gap: 8px;
}

.stRadio label {
    background: rgba(255, 255, 255, 0.74);
    border: 1px solid rgba(15, 23, 42, 0.08);
    border-radius: 999px;
    padding: 5px 10px;
}

textarea,
input,
[data-baseweb="select"] > div {
    border-radius: 10px !important;
    border-color: rgba(15, 23, 42, 0.12) !important;
}

[data-testid="stDataFrame"] {
    border: 1px solid rgba(15, 23, 42, 0.08);
    border-radius: 10px;
    overflow: hidden;
    box-shadow: 0 12px 26px rgba(15, 23, 42, 0.04);
}

div[data-testid="stExpander"] {
    border-radius: 10px;
    border-color: rgba(15, 23, 42, 0.1);
    background: rgba(255, 255, 255, 0.72);
}

@media (max-width: 760px) {
    .cc-hero-grid {
        grid-template-columns: 1fr;
    }
    .cc-hero-metrics {
        grid-template-columns: 1fr;
    }
    .cc-hero h1 {
        font-size: 2.4rem;
    }
    .cc-compass-mark {
        display: none;
    }
}
</style>
"""


ROLE_STATUS = [
    {
        "owner": "Carlo",
        "lane": "Architecture",
        "status": "In progress",
        "summary": "Supervisor workflow, shared state, inter-agent handoffs, and UI integration are wired.",
    },
    {
        "owner": "Nhi",
        "lane": "RAG/Data",
        "status": "Started",
        "summary": "Local job-posting evidence and retrieval scoring now feed the Market Demand Agent.",
    },
    {
        "owner": "TM3",
        "lane": "Agent Logic",
        "status": "Started",
        "summary": "Prompts, JSON contracts, validation, OpenAI hook, and fallbacks are scaffolded.",
    },
    {
        "owner": "TM4",
        "lane": "Frontend/Demo",
        "status": "Active now",
        "summary": "Streamlit presenter flow, demo script, and screenshot checklist are being prepared.",
    },
    {
        "owner": "TM5",
        "lane": "Evaluation/Deployment",
        "status": "Started",
        "summary": "Docker, evaluation plan, ethics notes, and deployment docs are now scaffolded.",
    },
]


DEMO_CUES = {
    "Dashboard": "Lead with the value proposition: CareerCompass turns a resume, coursework, and a target role into a coordinated career strategy.",
    "Roadmap": "Explain that recommendations become a time-bound learning plan, not generic advice.",
    "Resume": "Show keyword targets, selectable improvements, editable draft, and DOCX export.",
    "Interview": "Use Salesforce or another Bay Area company, generate scenario questions, answer one, then show scoring and sample answer.",
    "Final Report": "Call out that the supervisor synthesizes every specialist output into one student-facing strategy report.",
    "Action Checklist": "Close with practical next steps: certifications, missing evidence, and portfolio proof.",
    "Rubric Evidence": "Use this as the grading bridge: each project requirement maps to a feature, document, or verification artifact.",
}

DEMO_VIEW_PARAMS = {
    "dashboard": "Dashboard",
    "roadmap": "Roadmap",
    "resume": "Resume",
    "interview": "Interview",
    "report": "Final Report",
    "checklist": "Action Checklist",
    "rubric": "Rubric Evidence",
}

RUBRIC_EVIDENCE = [
    {
        "Rubric area": "Agent Architecture & Implementation",
        "Weight": "35%",
        "Status": "Ready",
        "Evidence": "LangGraph supervisor workflow, five specialist agents, synthesis node, typed shared state, structured handoffs, model fallbacks, and root workflow tests.",
    },
    {
        "Rubric area": "Business Value & Relevance",
        "Weight": "20%",
        "Status": "Ready",
        "Evidence": "Student and career-center use case with resume intake, target role/location/timeline, skill gaps, roadmap, resume draft, interview practice, and strategy report.",
    },
    {
        "Rubric area": "Container & Deployment Strategy",
        "Weight": "15%",
        "Status": "Ready for smoke test",
        "Evidence": "Dockerfile, .dockerignore, deployment guide, deterministic fallback mode, and optional OpenAI environment variables.",
    },
    {
        "Rubric area": "Ethical Consideration",
        "Weight": "15%",
        "Status": "Ready",
        "Evidence": "Visible Responsible AI audit, advisor-review warning, confidence scores, no intentional resume persistence, evidence grounding, and ethics documentation.",
    },
    {
        "Rubric area": "Documentation & Presentation",
        "Weight": "15%",
        "Status": "Ready for final media",
        "Evidence": "README, architecture handoff, demo script, screenshot checklist, evaluation docs, deployment docs, and this rubric evidence table.",
    },
]

RESPONSIBLE_AI_CHECKS = [
    {
        "Control": "Resume privacy",
        "Current implementation": "Resume text is used in the Streamlit session and is not intentionally persisted by the MVP.",
        "Demo note": "Production should add explicit retention and deletion controls before storing resumes.",
    },
    {
        "Control": "Grounded recommendations",
        "Current implementation": "Market skills and gaps are tied to retrieved job-posting evidence and student-provided resume/coursework signals.",
        "Demo note": "Current local corpus is directional, not a full labor-market sample.",
    },
    {
        "Control": "Human oversight",
        "Current implementation": "The final report frames output as decision support and recommends career-advisor review.",
        "Demo note": "Avoid promising guaranteed job matches or definitive hiring outcomes.",
    },
    {
        "Control": "Resume authenticity",
        "Current implementation": "Resume suggestions preserve evidence and warn against exaggerated claims.",
        "Demo note": "Students should edit generated bullets before submitting applications.",
    },
]

RAG_READINESS_ROWS = [
    {
        "Layer": "Current RAG",
        "Status": "Implemented",
        "Evidence": "Deterministic lexical retrieval over local job postings returns scored sources and evidence summaries.",
    },
    {
        "Layer": "Vector-search upgrade",
        "Status": "Documented next step",
        "Evidence": "ChromaDB/Kaggle LinkedIn path is documented as the next RAG lane while preserving the current posting shape.",
    },
    {
        "Layer": "Bias limitation",
        "Status": "Disclosed",
        "Evidence": "The app and docs describe the local corpus as directional and recommend advisor review.",
    },
]

INTERVIEW_PROGRESS_LABELS = {
    "not_started": "Not started",
    "questions_ready": "Questions ready",
    "practice_started": "Practice started",
    "feedback_received": "Feedback received",
    "ready_to_apply": "Ready to apply",
}

INTERVIEW_PROGRESS_ORDER = {
    "not_started": 0,
    "questions_ready": 1,
    "practice_started": 2,
    "feedback_received": 3,
    "ready_to_apply": 4,
}

INTERVIEW_STATUS_HELP = (
    "Generate questions, answer one, and evaluate your response. "
    "A 4/5 or higher marks interview prep as ready."
)


st.set_page_config(
    page_title="CareerCompass",
    page_icon="🧭",
    layout="wide",
    initial_sidebar_state="expanded",
)


def apply_visual_style() -> None:
    st.markdown(ACCENT_CSS, unsafe_allow_html=True)


def initialize_state() -> None:
    defaults = {
        "analysis": None,
        "latency_seconds": None,
        "active_view": "Dashboard",
        "last_resume_upload": None,
        "resume_upload_notice": None,
        "resume_text_area": SAMPLE_RESUME,
        "target_role_choice": "Business Analyst",
        "custom_target_role": "",
        "target_location_choice": "San Francisco Bay Area",
        "custom_target_location": "",
        "coursework_selection": SAMPLE_COURSEWORK,
        "additional_coursework": "",
        "target_job_company": "",
        "target_job_title": "",
        "target_job_url": "",
        "target_job_description": "",
        "resume_target_job_company": "",
        "resume_target_job_title": "",
        "resume_target_job_url": "",
        "resume_target_job_description": "",
        "resume_draft": "",
        "selected_resume_template": "Use my existing resume structure",
        "custom_resume_template": "",
        "resume_profile_resume_editor": "",
        "resume_profile_source_snapshot": "",
        "safe_edits_only": False,
        "resume_regeneration_mode": "Standard",
        "resume_export_reviewed": False,
        "tailoring_error": "",
        "tailored_resume_history": [],
        "tailoring_summary": None,
        "tailoring_change_rows": [],
        "practice_questions": None,
        "interview_evaluation": None,
        "interview_progress_stage": "not_started",
        "interview_latest_score": None,
        "interview_answer_draft": "",
        "demo_mode": False,
        "analysis_validation_error": "",
        "resume_upload_metadata": None,
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def reset_demo_profile() -> None:
    st.session_state.analysis = None
    st.session_state.latency_seconds = None
    st.session_state.active_view = "Dashboard"
    st.session_state.last_resume_upload = None
    st.session_state.resume_upload_notice = "Sample MIS graduate profile restored."
    st.session_state.resume_text_area = SAMPLE_RESUME
    st.session_state.target_role_choice = "Business Analyst"
    st.session_state.custom_target_role = ""
    st.session_state.target_location_choice = "San Francisco Bay Area"
    st.session_state.custom_target_location = ""
    st.session_state.coursework_selection = SAMPLE_COURSEWORK
    st.session_state.additional_coursework = ""
    st.session_state.target_job_company = ""
    st.session_state.target_job_title = ""
    st.session_state.target_job_url = ""
    st.session_state.target_job_description = ""
    st.session_state.resume_target_job_company = ""
    st.session_state.resume_target_job_title = ""
    st.session_state.resume_target_job_url = ""
    st.session_state.resume_target_job_description = ""
    st.session_state.resume_draft = ""
    st.session_state.selected_resume_template = "Use my existing resume structure"
    st.session_state.custom_resume_template = ""
    st.session_state.resume_profile_resume_editor = SAMPLE_RESUME
    st.session_state.resume_profile_source_snapshot = SAMPLE_RESUME
    st.session_state.safe_edits_only = False
    st.session_state.resume_regeneration_mode = "Standard"
    st.session_state.resume_export_reviewed = False
    st.session_state.tailoring_error = ""
    st.session_state.tailored_resume_history = []
    st.session_state.tailoring_summary = None
    st.session_state.tailoring_change_rows = []
    st.session_state.practice_questions = None
    st.session_state.interview_evaluation = None
    st.session_state.interview_progress_stage = "not_started"
    st.session_state.interview_latest_score = None
    st.session_state.interview_answer_draft = ""
    st.session_state.demo_autorun_view = None


def update_saved_resume_from_editor() -> None:
    edited_resume = st.session_state.get("resume_profile_resume_editor", "")
    st.session_state.resume_text_area = edited_resume
    st.session_state.resume_profile_source_snapshot = edited_resume.strip()


def clear_resume_job_post() -> None:
    st.session_state.resume_target_job_company = ""
    st.session_state.resume_target_job_title = ""
    st.session_state.resume_target_job_url = ""
    st.session_state.resume_target_job_description = ""
    st.session_state.resume_export_reviewed = False
    st.session_state.tailoring_error = ""
    st.session_state.tailoring_summary = None
    st.session_state.tailoring_change_rows = []


def upload_metadata(filename: str, text: str) -> dict:
    suffix = filename.rsplit(".", 1)[-1].upper() if "." in filename else "TEXT"
    words = re.findall(r"\b\w+\b", text)
    sections = list(parse_resume_sections(text).keys())[:6] if text.strip() else []
    return {
        "Filename": filename,
        "File type": suffix,
        "Characters": len(text),
        "Words": len(words),
        "Detected sections": ", ".join(sections) if sections else "No sections detected",
    }


def validate_resume_for_analysis(resume_text: str) -> str:
    if not resume_text.strip():
        return "Add a resume before running analysis."
    return ""


def validate_job_post_for_tailoring(job_post: str) -> str:
    words = re.findall(r"\b\w+\b", job_post)
    if not job_post.strip():
        return "Paste the job description you want to apply to before tailoring your resume."
    if len(words) < 35:
        return "Paste a fuller job posting so CareerCompass can compare responsibilities, qualifications, and keywords."
    return ""


def validate_interview_answer(answer: str) -> str:
    if not answer.strip():
        return "Write or paste an answer before requesting feedback."
    return ""


def reset_interview_progress() -> None:
    st.session_state.practice_questions = None
    st.session_state.interview_evaluation = None
    st.session_state.interview_progress_stage = "not_started"
    st.session_state.interview_latest_score = None
    st.session_state.interview_answer_draft = ""


def advance_interview_progress(stage: str, score: int | None = None) -> None:
    current = st.session_state.get("interview_progress_stage", "not_started")
    if INTERVIEW_PROGRESS_ORDER[stage] >= INTERVIEW_PROGRESS_ORDER.get(current, 0):
        st.session_state.interview_progress_stage = stage
    if score is not None:
        st.session_state.interview_latest_score = score


def interview_status_label() -> str:
    stage = st.session_state.get("interview_progress_stage", "not_started")
    return INTERVIEW_PROGRESS_LABELS.get(stage, INTERVIEW_PROGRESS_LABELS["not_started"])


def section_header(title: str, help_text: str, level: str = "subheader") -> None:
    if level == "header":
        st.header(title, help=help_text)
    else:
        st.subheader(title, help=help_text)


def explain_market_skill_row(row: pd.Series) -> str:
    skill = str(row.get("Skill", "this skill"))
    demand = str(row.get("Demand Signal", "")).lower()
    if "very high" in demand:
        return f"Make {skill} easy to find in your resume and prepare an interview example."
    if "high" in demand:
        return f"Add truthful {skill} evidence where it supports the job posting."
    if "medium" in demand:
        return f"Use {skill} as supporting evidence when it matches your experience."
    return f"Consider {skill} only if you have real resume or portfolio proof."


def target_job_from_session(prefix: str = "target_job") -> dict[str, str]:
    return {
        "company": st.session_state.get(f"{prefix}_company", "").strip(),
        "title": st.session_state.get(f"{prefix}_title", "").strip(),
        "url": st.session_state.get(f"{prefix}_url", "").strip(),
        "description": st.session_state.get(f"{prefix}_description", "").strip(),
    }


def target_job_to_text(target_job: dict[str, str]) -> str:
    parts = [
        target_job.get("title", ""),
        target_job.get("company", ""),
        target_job.get("url", ""),
        target_job.get("description", ""),
    ]
    return "\n".join(part for part in parts if part).strip()


def sync_resume_target_job_from_profile() -> None:
    st.session_state.resume_target_job_company = st.session_state.target_job_company
    st.session_state.resume_target_job_title = st.session_state.target_job_title
    st.session_state.resume_target_job_url = st.session_state.target_job_url
    st.session_state.resume_target_job_description = st.session_state.target_job_description


def sync_profile_target_job_from_resume() -> None:
    st.session_state.generated_resume_target_job = target_job_from_session("resume_target_job")


def sample_demo_inputs() -> dict:
    return {
        "resume_text": SAMPLE_RESUME,
        "target_role": "Business Analyst",
        "target_location": "San Francisco Bay Area",
        "timeline_days": 90,
        "coursework": SAMPLE_COURSEWORK,
        "target_job": {},
    }


def apply_demo_query_params() -> None:
    if st.query_params.get("demo_autorun") != "1":
        return

    view_param = st.query_params.get("demo_view", "dashboard")
    st.session_state.active_view = DEMO_VIEW_PARAMS.get(view_param, "Dashboard")

    if st.session_state.get("demo_autorun_view") == view_param and st.session_state.analysis:
        return

    started_at = time.perf_counter()
    st.session_state.analysis = run_career_analysis(sample_demo_inputs())
    st.session_state.latency_seconds = time.perf_counter() - started_at
    reset_interview_progress()
    st.session_state.demo_autorun_view = view_param


def render_sidebar() -> None:
    st.sidebar.markdown(
        """
        <div class="cc-sidebar-brand">
            <div class="cc-sidebar-compass" aria-hidden="true"></div>
            <div>
                <strong>CareerCompass</strong>
                <span>Navigate your career</span>
            </div>
        </div>
        <div class="cc-sidebar-label">Command Center</div>
        <div class="cc-sidebar-route">
            <div class="cc-sidebar-step"><strong>Profile Scan</strong><span>Resume, role, location, coursework</span></div>
            <div class="cc-sidebar-step"><strong>Market Signal</strong><span>Role evidence and skill demand</span></div>
            <div class="cc-sidebar-step"><strong>Resume Targeting</strong><span>Skill proof and job-specific drafts</span></div>
            <div class="cc-sidebar-step"><strong>Interview Prep</strong><span>Scenario practice and answer feedback</span></div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.sidebar.divider()
    st.sidebar.toggle(
        "Demo / Instructor mode",
        key="demo_mode",
        help="Shows sample reset controls and technical evidence used for class demos.",
    )


def render_role_status_panel() -> None:
    columns = st.columns(len(ROLE_STATUS))
    for column, item in zip(columns, ROLE_STATUS):
        with column:
            with st.container(border=True):
                st.markdown(f"**{item['owner']} - {item['lane']}**")
                st.caption(item["status"].upper())
                st.write(item["summary"])


def render_presenter_cue(view_name: str) -> None:
    if not st.session_state.demo_mode:
        return

    cue = DEMO_CUES.get(view_name)
    if not cue:
        return

    st.markdown(
        f"""
        <div class="cc-demo-cue">
            <strong>Presenter cue:</strong> {cue}
        </div>
        """,
        unsafe_allow_html=True,
    )


def extract_uploaded_resume(uploaded_file) -> tuple[str | None, str]:
    suffix = uploaded_file.name.lower().rsplit(".", 1)[-1]
    raw = uploaded_file.getvalue()

    if suffix in {"txt", "md"}:
        return raw.decode("utf-8", errors="ignore"), "Resume text loaded from uploaded text file."

    if suffix == "pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            return None, "PDF upload is wired in, but install `pypdf` to extract PDF text."

        reader = PdfReader(uploaded_file)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        return text.strip(), "Resume text extracted from uploaded PDF."

    if suffix == "docx":
        try:
            from docx import Document
        except ImportError:
            return None, "DOCX upload is wired in, but install `python-docx` to extract DOCX text."

        document = Document(uploaded_file)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        return text.strip(), "Resume text extracted from uploaded DOCX."

    return None, "Unsupported file type. Use TXT, PDF, or DOCX."


def render_inputs() -> dict:
    st.header("Build Your Career Profile")
    st.caption(
        "Paste your resume text or upload a file. We'll use only this session unless you choose to save it."
    )

    if st.session_state.demo_mode:
        reset_col, note_col = st.columns([0.35, 0.65])
        with reset_col:
            if st.button("Reset sample demo profile"):
                reset_demo_profile()
                st.rerun()
        with note_col:
            st.caption(
                "Demo mode: reset before recording so the walkthrough starts from a predictable sample profile."
            )

    left, right = st.columns([1.25, 1])

    with left:
        uploaded_resume = st.file_uploader(
            "Upload resume",
            type=["txt", "md", "pdf", "docx"],
            help="TXT works immediately. PDF and DOCX extraction are supported when optional parser packages are installed.",
        )
        if uploaded_resume and uploaded_resume.name != st.session_state.last_resume_upload:
            extracted_text, notice = extract_uploaded_resume(uploaded_resume)
            st.session_state.last_resume_upload = uploaded_resume.name
            st.session_state.resume_upload_notice = notice
            if extracted_text:
                st.session_state.resume_text_area = extracted_text
                st.session_state.resume_upload_metadata = upload_metadata(uploaded_resume.name, extracted_text)

        if st.session_state.resume_upload_notice:
            st.info(st.session_state.resume_upload_notice)
        if st.session_state.resume_upload_metadata:
            with st.expander("Resume parsing confirmation", expanded=True):
                st.dataframe(
                    pd.DataFrame([st.session_state.resume_upload_metadata]),
                    use_container_width=True,
                    hide_index=True,
                )

        resume_text = st.text_area(
            "Resume text",
            height=300,
            key="resume_text_area",
            help="You can edit uploaded text here before running the analysis.",
        )
        if resume_text.strip():
            pasted_metadata = upload_metadata("Pasted or edited resume text", resume_text)
            st.success(
                "Resume added: "
                f"{pasted_metadata['Words']} words, "
                f"{pasted_metadata['Characters']} characters, "
                f"detected sections: {pasted_metadata['Detected sections']}."
            )
            with st.expander("Current resume text confirmation", expanded=False):
                st.dataframe(
                    pd.DataFrame([pasted_metadata]),
                    use_container_width=True,
                    hide_index=True,
                )

    with right:
        role_options = available_target_roles() + ["Custom role"]
        if st.session_state.target_role_choice not in role_options:
            st.session_state.target_role_choice = role_options[0]

        target_role_choice = st.selectbox(
            "Target role",
            options=role_options,
            key="target_role_choice",
            help="Choose a role from the current job-posting dataset, or use Custom role for another target.",
        )
        if target_role_choice == "Custom role":
            target_role = st.text_input(
                "Custom target role",
                key="custom_target_role",
                placeholder="Example: Product Analyst, Data Engineer, UX Researcher",
            ).strip()
        else:
            target_role = target_role_choice

        location_options = available_locations() + ["Custom location"]
        if st.session_state.target_location_choice not in location_options:
            st.session_state.target_location_choice = location_options[0]

        target_location_choice = st.selectbox(
            "Target location",
            options=location_options,
            key="target_location_choice",
            help="Locations come from the current local job-posting dataset.",
        )
        if target_location_choice == "Custom location":
            target_location = st.text_input(
                "Custom target location",
                key="custom_target_location",
                placeholder="Example: Los Angeles, Remote, New York",
            ).strip()
        else:
            target_location = target_location_choice

        timeline_days = st.slider(
            "Timeline (days until target hire date)",
            min_value=30,
            max_value=180,
            value=90,
            step=30,
        )
        coursework_options = available_coursework_options()
        selected_defaults = [
            option for option in st.session_state.coursework_selection if option in coursework_options
        ]
        if selected_defaults != st.session_state.coursework_selection:
            st.session_state.coursework_selection = selected_defaults

        coursework = st.multiselect(
            "Relevant coursework",
            options=coursework_options,
            key="coursework_selection",
            help="Options are expanded from the skills represented in the local job-posting dataset.",
        )
        additional_coursework = st.text_area(
            "Additional coursework or training",
            key="additional_coursework",
            height=92,
            placeholder="Add courses, certifications, bootcamps, or workshops separated by commas or new lines.",
        )
        coursework = merge_coursework(coursework, additional_coursework)

    st.subheader("Target Job Posting")
    st.caption(
        "Paste the job description you want to apply to. We'll compare it against your resume and show exactly what changed."
    )
    job_meta_col1, job_meta_col2 = st.columns([1, 1])
    with job_meta_col1:
        st.text_input("Company", key="target_job_company", placeholder="Example: Jamba Juice")
        st.text_input("Job title", key="target_job_title", placeholder="Example: Product Marketing Associate")
    with job_meta_col2:
        st.text_input("Posting URL (optional)", key="target_job_url", placeholder="https://...")
    st.text_area(
        "Full job description",
        key="target_job_description",
        height=170,
        placeholder="Paste responsibilities, qualifications, preferred skills, and company context here.",
    )

    target_job = target_job_from_session("target_job")
    if target_job["description"]:
        word_count = len(re.findall(r"\b\w+\b", target_job["description"]))
        st.caption(f"Target job saved for this session: {word_count} words from the pasted posting.")

    return {
        "resume_text": resume_text,
        "target_role": target_role or "Business Analyst",
        "target_location": target_location or "San Francisco Bay Area",
        "timeline_days": timeline_days,
        "coursework": coursework,
        "target_job": target_job,
    }


def merge_coursework(selected: list[str], additional_text: str) -> list[str]:
    additional = [
        chunk.strip()
        for chunk in additional_text.replace("\n", ",").split(",")
        if chunk.strip()
    ]
    merged = []
    for item in [*selected, *additional]:
        if item not in merged:
            merged.append(item)
    return merged


def render_progress() -> None:
    agents = [
        "Parsing resume evidence",
        "Reading target job post",
        "Matching skills and gaps",
        "Building route plan",
        "Preparing resume and interview guidance",
        "Finalizing navigation brief",
    ]

    progress = st.progress(0)
    status = st.empty()

    for index, agent in enumerate(agents, start=1):
        status.write(f"{agent}...")
        progress.progress(index / len(agents))
        time.sleep(0.18)

    status.success("Career strategy generated.")


def render_tailoring_progress() -> None:
    steps = [
        "Parsing job post",
        "Matching resume evidence",
        "Drafting safe edits",
        "Checking unsupported claims",
    ]
    progress = st.progress(0)
    status = st.empty()
    for index, step in enumerate(steps, start=1):
        status.write(f"{step}...")
        progress.progress(index / len(steps))
        time.sleep(0.12)
    status.success("Tailored resume generated.")


def friendly_tailoring_error(error: Exception) -> str:
    return (
        "Resume tailoring hit a recoverable issue. Your saved resume and job post are still here; "
        "review the fields and try Generate tailored resume again."
    )


def resume_export_allowed(change_rows: list[dict], reviewed: bool) -> bool:
    return not any(row.get("Safe To Add") != "Yes" for row in change_rows) or reviewed


def render_dashboard(analysis: dict) -> None:
    section_header(
        "Career Readiness Console",
        "A quick scan of your role fit, highest-priority gaps, resume coverage, and interview-prep progress.",
        level="header",
    )
    st.markdown(
        '<div class="cc-page-signal">Live scan of role fit, priority gaps, market demand, and next moves.</div>',
        unsafe_allow_html=True,
    )

    c1, c2, c3, c4 = st.columns(4)
    c1.metric(
        "Role fit",
        f"{analysis['match_percentage']}%",
        help="How closely your profile currently lines up with this target role and location.",
    )
    c2.metric(
        "Top gaps",
        analysis["gap_counts"]["high"],
        help="High-priority gaps CareerCompass recommends addressing before applying.",
    )
    c3.metric(
        "Resume match",
        f"{analysis['keyword_coverage']}%",
        help="How much target skill and evidence coverage CareerCompass sees in your resume.",
    )
    c4.metric(
        "Interview status",
        interview_status_label(),
        help=INTERVIEW_STATUS_HELP,
    )

    section_header(
        "Recommended Next Moves",
        "The highest-leverage actions to take next, ordered by timing and role impact.",
    )
    next_actions = pd.DataFrame(analysis["next_actions"])
    st.dataframe(
        next_actions,
        use_container_width=True,
        hide_index=True,
        column_config={
            "Priority": st.column_config.TextColumn(
                "Priority",
                help="When this action should happen in your application prep.",
            ),
            "Action": st.column_config.TextColumn(
                "Action",
                help="The concrete next step CareerCompass recommends.",
            ),
            "Why it matters": st.column_config.TextColumn(
                "Why it matters",
                help="Why this action improves your readiness for the target role.",
            ),
        },
    )

    section_header(
        "Market Demand Signal",
        "Skills and proof points that appear in retrieved postings or the pasted job description. Use this to decide what belongs in your resume, portfolio, and interview stories.",
    )
    skills = pd.DataFrame(analysis["market_skills"])
    if not skills.empty:
        skills["What this means"] = skills.apply(explain_market_skill_row, axis=1)
    st.dataframe(
        skills,
        use_container_width=True,
        hide_index=True,
        column_config={
            "Skill": st.column_config.TextColumn(
                "Skill",
                help="A role-relevant skill or proof point to make visible.",
            ),
            "Demand Signal": st.column_config.TextColumn(
                "Demand Signal",
                help="How often this skill appears in the retrieved evidence for the target role.",
            ),
            "Evidence": st.column_config.TextColumn(
                "Evidence",
                help="Where CareerCompass saw this demand signal.",
            ),
            "What this means": st.column_config.TextColumn(
                "What this means",
                help="Plain-language guidance for how to use this signal.",
            ),
        },
    )
    retrieved_postings = analysis.get("retrieved_job_postings", [])
    if retrieved_postings:
        with st.expander("Retrieved job-posting evidence"):
            evidence_rows = [
                {
                    "Company": posting["company"],
                    "Role": posting["role"],
                    "Location": posting["location"],
                    "Score": posting["retrieval_score"],
                    "Evidence": posting["evidence_summary"],
                }
                for posting in retrieved_postings
            ]
            st.dataframe(pd.DataFrame(evidence_rows), use_container_width=True, hide_index=True)

    section_header(
        "Skill Gap Radar",
        "The skills where your current resume/coursework evidence looks weakest compared with the target role. Start with High severity rows.",
    )
    gaps = pd.DataFrame(analysis["gap_report"])
    st.dataframe(
        gaps,
        use_container_width=True,
        hide_index=True,
        column_config={
            "Skill": st.column_config.TextColumn(
                "Skill",
                help="The missing or underdeveloped skill area.",
            ),
            "Current Evidence": st.column_config.TextColumn(
                "Current Evidence",
                help="What CareerCompass can currently find in your resume and coursework.",
            ),
            "Severity": st.column_config.TextColumn(
                "Severity",
                help="How urgently this gap should be addressed for the target role.",
            ),
            "Recommendation": st.column_config.TextColumn(
                "Recommendation",
                help="A practical way to reduce this gap.",
            ),
            "First Step": st.column_config.TextColumn(
                "First Step",
                help="The smallest useful action to begin closing this gap.",
            ),
            "Resume Proof": st.column_config.TextColumn(
                "Resume Proof",
                help="The kind of truthful resume evidence to create or clarify.",
            ),
        },
    )

    section_header(
        "Suggested Certifications and Short Courses",
        "Optional learning signals that can help fill gaps. Prioritize free or fast options that support your target role evidence.",
    )
    st.dataframe(
        pd.DataFrame(analysis["certification_recommendations"]),
        use_container_width=True,
        hide_index=True,
        column_config={
            "Area": st.column_config.TextColumn(
                "Area",
                help="The gap or capability this resource supports.",
            ),
            "Recommendation": st.column_config.TextColumn(
                "Recommendation",
                help="A suggested course, credential, or learning artifact.",
            ),
            "Priority": st.column_config.TextColumn(
                "Priority",
                help="How useful this resource is for the current target.",
            ),
            "Why": st.column_config.TextColumn(
                "Why",
                help="Why this resource may improve your application evidence.",
            ),
        },
    )

    section_header(
        "Start Closing a Gap",
        "Pick one gap and turn it into a small project, resume bullet, or interview story.",
    )
    selected_gap = st.selectbox(
        "Choose a skill gap to work on",
        options=[gap["Skill"] for gap in analysis["gap_report"]],
        help="CareerCompass can turn each recommendation into a first project plan.",
    )
    deep_dive = analysis["gap_deep_dives"][selected_gap]

    with st.container(border=True):
        st.markdown(f"**Why {selected_gap} matters**")
        st.write(deep_dive["why_it_matters"])

        case_col, step_col = st.columns([1, 1])
        with case_col:
            st.markdown("**Starter business cases**")
            for idea in deep_dive["starter_business_cases"]:
                st.checkbox(idea, value=False, key=f"case-{selected_gap}-{idea}")

        with step_col:
            st.markdown("**First steps**")
            for step_number, step in enumerate(deep_dive["first_steps"], start=1):
                st.write(f"{step_number}. {step}")

        st.markdown("**What to research next**")
        st.write(", ".join(deep_dive["search_terms"]))
        st.markdown("**Resume proof to create**")
        st.success(deep_dive["resume_proof"])

    if st.session_state.demo_mode:
        with st.expander("Technical demo notes for the class report"):
            st.write(
                "Latency and evaluation evidence are useful for the project rubric, but less useful as front-and-center student metrics."
            )
            st.metric("End-to-end latency", f"{st.session_state.latency_seconds:.1f}s")
            st.subheader("Supervisor agent trace")
            st.write(" -> ".join(analysis.get("agent_trace", [])))
            st.subheader("Agent coordination evidence")
            st.dataframe(
                pd.DataFrame(build_agent_coordination_rows(analysis)),
                use_container_width=True,
                hide_index=True,
            )

            handoffs = analysis.get("agent_handoffs", [])
            if handoffs:
                st.subheader("Raw inter-agent handoffs")
                st.dataframe(pd.DataFrame(handoffs), use_container_width=True, hide_index=True)

            confidence_scores = analysis.get("confidence_scores", {})
            if confidence_scores:
                st.subheader("Confidence scores")
                confidence_rows = [
                    {"Signal": key.replace("_", " ").title(), "Score": value}
                    for key, value in confidence_scores.items()
                ]
                st.dataframe(pd.DataFrame(confidence_rows), use_container_width=True, hide_index=True)

            st.subheader("Evaluation evidence")
            st.dataframe(
                pd.DataFrame(analysis["evaluation_metrics"]),
                use_container_width=True,
                hide_index=True,
            )

    render_responsible_ai_audit(analysis)


def render_roadmap(analysis: dict) -> None:
    st.header("90-Day Route Plan")
    st.markdown(
        '<div class="cc-page-signal">Milestones for turning gaps into visible career proof.</div>',
        unsafe_allow_html=True,
    )
    cols = st.columns(3)
    for col, phase in zip(cols, analysis["learning_roadmap"]):
        with col:
            st.subheader(phase["period"])
            st.write(phase["goal"])
            for task in phase["tasks"]:
                st.checkbox(task, value=False, key=f"{phase['period']}-{task}")
            st.caption(f"Resource relevance: {phase['resource_relevance']}/5")

    st.subheader("Portfolio Proof Waypoints")
    for item in analysis["portfolio_ideas"]:
        st.checkbox(item, value=False, key=f"portfolio-{item}")


def render_resume(analysis: dict) -> None:
    st.header("Resume Targeting Console")
    st.caption("Reuse the saved profile resume, paste a job post, map skill signals, generate a targeted draft, then export.")

    st.subheader("Resume Keyword Targets")
    st.write(
        "These are the words and phrases CareerCompass thinks should be represented somewhere in the resume."
    )
    keyword_df = pd.DataFrame(analysis["keyword_targets"])
    st.dataframe(keyword_df, use_container_width=True, hide_index=True)

    st.subheader("Resume Evidence Reminders")
    st.caption("Use these while reviewing the tailored draft for this job post.")
    for item in analysis["resume_checklist"]:
        st.checkbox(item, value=False, key=f"resume-check-{item}")

    saved_resume = st.session_state.resume_text_area.strip()
    if st.session_state.resume_profile_source_snapshot != saved_resume:
        st.session_state.resume_profile_resume_editor = saved_resume
        st.session_state.resume_profile_source_snapshot = saved_resume
    if not st.session_state.resume_target_job_description and st.session_state.target_job_description:
        sync_resume_target_job_from_profile()

    st.subheader("Tailor Your Resume to a Specific Job")
    st.caption(
        "Paste the job description here. We'll use your saved resume and the job requirements to rebuild a targeted version of your resume."
    )
    with st.expander("Review or replace the saved profile resume"):
        st.write("This is the resume from the Build your career profile step. Edit it here only if you want to replace the saved version for this application batch.")
        edited_saved_resume = st.text_area(
            "Saved profile resume",
            height=260,
            key="resume_profile_resume_editor",
        )
        if st.button("Update saved profile resume from this editor", on_click=update_saved_resume_from_editor):
            st.success("Saved profile resume updated for future tailored drafts.")

    if st.session_state.target_job_description:
        st.info("Using the target job posting saved in the profile step. Edit the fields below for this application.")

    job_col1, job_col2 = st.columns([1, 1])
    with job_col1:
        st.text_input("Company", key="resume_target_job_company")
        st.text_input("Job title", key="resume_target_job_title")
    with job_col2:
        st.text_input("Posting URL (optional)", key="resume_target_job_url")
        if st.button("Sync profile job to resume tab", on_click=sync_resume_target_job_from_profile):
            st.success("Profile target job copied into the resume builder.")
    job_description = st.text_area(
        "Paste the job post you're applying to",
        key="resume_target_job_description",
        height=240,
        placeholder="Paste the full job description, responsibilities, qualifications, and preferred skills here.",
    )
    target_job = target_job_from_session("resume_target_job")
    job_post = target_job_to_text(target_job)
    job_keywords = derive_job_post_keywords(job_post, analysis)
    if job_keywords:
        st.markdown("**Job-specific keyword signals**")
        st.dataframe(pd.DataFrame(job_keywords), use_container_width=True, hide_index=True)
    else:
        st.info("Paste a job post to add job-specific keywords to the existing CareerCompass resume targets.")

    st.subheader("Signal Map")
    st.caption("Use this to see what CareerCompass found before generating the tailored resume.")
    if job_post:
        evidence_map = build_skill_evidence_map(
            analysis=analysis,
            resume_text=st.session_state.resume_text_area,
            coursework=merge_coursework(
                st.session_state.coursework_selection,
                st.session_state.additional_coursework,
            ),
            job_post=job_post,
        )
        st.dataframe(pd.DataFrame(evidence_map), use_container_width=True, hide_index=True)
    else:
        st.info("Paste the full posting, not just the title, so we can compare responsibilities, qualifications, and keywords.")

    st.subheader("Choose a Resume Format")
    st.write("The selected format will restructure your saved resume for the pasted job post, not just apply a generic template.")
    format_options = [
        "Use my existing resume structure",
        "Project-forward resume",
        "Experience-forward resume",
        "Skills matrix resume",
        "Custom student-supplied template",
    ]
    if st.session_state.selected_resume_template not in format_options:
        st.session_state.selected_resume_template = "Use my existing resume structure"
    format_name = st.selectbox(
        "Resume format",
        options=format_options,
        key="selected_resume_template",
        help="Choose how the tailored resume should be organized for this job post.",
    )
    template_name = resume_format_template_key(format_name)

    custom_template = ""
    if format_name == "Custom student-supplied template":
        custom_template = st.text_area(
            "Paste your template or section order",
            value=st.session_state.custom_resume_template,
            height=180,
            key="custom_resume_template",
            placeholder="Example: Header, Summary, Education, Technical Skills, Projects, Experience, Certifications",
        )
        template = {
            "best_for": "A custom resume structure supplied by the student.",
            "use_when": "A professor, advisor, recruiter, or target industry expects a specific section order.",
            "avoid_when": "You are unsure whether the custom structure is ATS-friendly or evidence-first.",
            "sections": [line.strip() for line in custom_template.splitlines() if line.strip()],
            "preview": build_resume_draft(analysis, custom_template=custom_template),
        }
    else:
        template = resume_format_metadata(format_name, analysis, template_name)

    st.markdown("**Format tradeoffs**")
    card_cols = st.columns(2)
    for index, option in enumerate(format_options):
        option_template = (
            {"best_for": "A custom structure supplied by the student.", "use_when": "You already have a required template.", "avoid_when": "You do not have a clear section order.", "sections": ["Custom order"]}
            if option == "Custom student-supplied template"
            else resume_format_metadata(option, analysis, resume_format_template_key(option))
        )
        with card_cols[index % 2]:
            with st.container(border=True):
                st.markdown(f"**{option}**")
                st.caption(f"Best for: {option_template['best_for']}")
                st.write(f"Use when: {option_template['use_when']}")
                st.write(f"Avoid when: {option_template['avoid_when']}")
                st.caption("Section order: " + " -> ".join(option_template["sections"][:6]))

    t1, t2 = st.columns([0.9, 1.1])
    with t1:
        st.markdown(f"**Selected best for:** {template['best_for']}")
        st.markdown(f"**Use when:** {template['use_when']}")
        st.markdown(f"**Avoid when:** {template['avoid_when']}")
        st.markdown("**Recommended sections**")
        for section in template["sections"] or ["Custom section order will appear here."]:
            st.write(f"- {section}")
    with t2:
        preview = build_resume_draft(analysis, template_name, custom_template)
        st.text_area("Section-order preview", value=preview, height=260)
        if st.button("Save structure preview to editable resume"):
            st.session_state.resume_draft = st.session_state.resume_text_area.strip() or preview
            st.success("Structure preview copied into the editable resume draft.")

    st.subheader("Select Improvements to Apply")
    selected_suggestions = []

    for index, suggestion in enumerate(analysis["resume_recommendations"], start=1):
        with st.container(border=True):
            selected = st.checkbox(
                f"Apply suggestion {index}",
                value=False,
                key=f"resume-suggestion-{index}",
            )
            if selected:
                selected_suggestions.append(suggestion)
            before, after = st.columns(2)
            before.caption("Original")
            before.write(suggestion["before"])
            after.caption("Optimized")
            after.write(suggestion["after"])
            st.caption("Keywords added: " + ", ".join(suggestion["keywords_added"]))

    apply_col, reset_col = st.columns([1, 1])
    with apply_col:
        if st.button("Apply selected suggestions", type="primary"):
            draft = st.session_state.resume_draft or build_resume_draft(
                analysis,
                template_name,
                custom_template,
            )
            st.session_state.resume_draft = apply_resume_suggestions(draft, selected_suggestions)
            st.success("Selected suggestions applied to the editable resume draft.")
    with reset_col:
        if st.button("Reset draft from selected template"):
            st.session_state.resume_draft = build_resume_draft(
                analysis,
                template_name,
                custom_template,
            )

    st.subheader("Generate Job-Specific Resume")
    st.caption(
        "Review every suggested edit before submitting. CareerCompass will flag unsupported claims, but you are responsible for final accuracy."
    )
    mode_col, safety_col = st.columns([1, 1])
    with mode_col:
        regeneration_mode = st.selectbox(
            "Regenerate mode",
            options=["Standard", "Stricter evidence", "More concise", "ATS-focused"],
            key="resume_regeneration_mode",
            help="Use stricter evidence for conservative drafts, concise for shorter language, or ATS-focused for keyword-forward structure.",
        )
    with safety_col:
        safe_edits_only = st.toggle(
            "Safe edits only",
            key="safe_edits_only",
            help="When on, only strongly evidenced edits can enter the resume body. Missing claims always stay out.",
        )

    generate_col, clear_col = st.columns([1, 1])
    with generate_col:
        if st.button("Generate tailored resume", type="primary"):
            resume_error = validate_resume_for_analysis(st.session_state.resume_text_area)
            job_error = validate_job_post_for_tailoring(job_post)
            if resume_error:
                st.warning(resume_error)
            elif job_error:
                st.warning(job_error)
            else:
                try:
                    st.session_state.generated_resume_target_job = target_job
                    st.session_state.resume_export_reviewed = False
                    st.session_state.tailoring_error = ""
                    render_tailoring_progress()
                    coursework = merge_coursework(
                        st.session_state.coursework_selection,
                        st.session_state.additional_coursework,
                    )
                    change_rows = build_tailoring_change_summary(
                        saved_resume=st.session_state.resume_text_area,
                        analysis=analysis,
                        job_post=job_post,
                        selected_suggestions=selected_suggestions,
                        coursework=coursework,
                        safe_edits_only=safe_edits_only or regeneration_mode == "Stricter evidence",
                    )
                    tailored = build_tailored_resume_draft(
                        analysis=analysis,
                        saved_resume=st.session_state.resume_text_area,
                        job_post=job_post,
                        format_name=format_name,
                        custom_template=custom_template,
                        selected_suggestions=selected_suggestions,
                        coursework=coursework,
                        safe_edits_only=safe_edits_only or regeneration_mode == "Stricter evidence",
                        regeneration_mode=regeneration_mode,
                        target_job=target_job,
                    )
                    job_keywords = derive_job_post_keywords(job_post, analysis)
                    st.session_state.resume_draft = tailored
                    st.session_state.tailoring_change_rows = change_rows
                    st.session_state.tailoring_summary = summarize_tailoring_result(
                        change_rows,
                        job_keywords,
                    )
                    st.session_state.tailored_resume_history = [
                        {
                            "Company": target_job.get("company") or "Target company",
                            "Job": infer_job_title(job_post, analysis, target_job),
                            "Format": format_name,
                            "Mode": regeneration_mode,
                            "Keyword count": len(job_keywords),
                            "Draft": tailored,
                        },
                        *st.session_state.tailored_resume_history[:4],
                    ]
                    summary = st.session_state.tailoring_summary
                    st.success(
                        f"Tailored resume generated. Found {summary['keywords']} job keywords, "
                        f"revised {summary['revised_bullets']} supported bullets, "
                        f"and flagged {summary['missing_skills']} claims that need evidence."
                    )
                except Exception as error:
                    st.session_state.tailoring_error = friendly_tailoring_error(error)
                    st.error(st.session_state.tailoring_error)
    with clear_col:
        if st.button("Clear job post for next application", on_click=clear_resume_job_post):
            st.success("Job post cleared. Your saved profile resume is still available.")

    if st.session_state.tailored_resume_history:
        st.caption("Recent tailored resume runs")
        history_display = [
            {key: value for key, value in item.items() if key != "Draft"}
            for item in st.session_state.tailored_resume_history
        ]
        st.dataframe(pd.DataFrame(history_display), use_container_width=True, hide_index=True)
        history_options = list(range(len(st.session_state.tailored_resume_history)))
        selected_history = st.selectbox(
            "Reload a recent tailored draft",
            options=history_options,
            format_func=lambda index: (
                f"{st.session_state.tailored_resume_history[index]['Job']} - "
                f"{st.session_state.tailored_resume_history[index]['Company']}"
            ),
        )
        if st.button("Reload selected draft"):
            st.session_state.resume_draft = st.session_state.tailored_resume_history[selected_history]["Draft"]
            st.success("Selected tailored draft loaded into the editor.")

    if st.session_state.tailoring_change_rows:
        st.subheader("Review Changes Before Applying")
        st.dataframe(pd.DataFrame(st.session_state.tailoring_change_rows), use_container_width=True, hide_index=True)

    if not st.session_state.resume_draft:
        st.session_state.resume_draft = st.session_state.resume_text_area.strip() or build_resume_draft(analysis, template_name, custom_template)

    st.subheader("Editable Resume Draft")
    st.session_state.resume_draft = st.text_area(
        "Draft resume",
        value=st.session_state.resume_draft,
        height=520,
        help="Review every suggested edit before submitting. Unsupported claims stay in their own section.",
    )

    if st.session_state.tailoring_summary:
        unsupported_count = st.session_state.tailoring_summary.get("missing_skills", 0)
        if unsupported_count:
            st.warning(
                f"{unsupported_count} claims are flagged under DO NOT ADD UNTIL YOU HAVE EVIDENCE. "
                "Review them before exporting."
            )
            st.checkbox(
                "I reviewed the unsupported claims and understand they should not be added without evidence.",
                key="resume_export_reviewed",
            )
        export_allowed = resume_export_allowed(
            st.session_state.tailoring_change_rows,
            st.session_state.get("resume_export_reviewed", False),
        )
        download_col1, download_col2 = st.columns([1, 1])
        with download_col1:
            st.download_button(
                "Export resume as TXT",
                data=st.session_state.resume_draft.encode("utf-8"),
                file_name="careercompass_resume.txt",
                mime="text/plain",
                disabled=not export_allowed,
            )
        with download_col2:
            st.download_button(
                "Export resume as DOCX",
                data=build_docx_bytes(st.session_state.resume_draft),
                file_name="careercompass_resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                disabled=not export_allowed,
            )
    else:
        st.info("Generate a tailored resume to unlock export actions for this job.")


def render_interview(analysis: dict) -> None:
    st.header("Interview Prep Console")
    st.caption(
        "Practice against generated questions, a company-specific scenario, or a scenario the student already received."
    )
    advance_interview_progress("questions_ready")

    scenario_options = list(analysis["interview_scenarios"].keys()) + ["Custom scenario"]
    c1, c2 = st.columns([1, 1])
    with c1:
        preset = st.selectbox("Scenario preset", options=scenario_options)
    with c2:
        company = st.text_input("Company or organization", value="Target company")

    default_scenario = (
        ""
        if preset == "Custom scenario"
        else analysis["interview_scenarios"][preset]
    )
    scenario = st.text_area(
        "Scenario details",
        value=default_scenario,
        height=110,
        help="Paste the exact prompt or interview scenario if the company already provided one.",
    )

    if st.button("Generate scenario questions"):
        st.session_state.practice_questions = generate_interview_questions(
            analysis["target_role"],
            company,
            scenario,
        )
        st.session_state.interview_evaluation = None
        st.session_state.interview_answer_draft = ""
        advance_interview_progress("questions_ready")

    questions = st.session_state.practice_questions or analysis["interview_questions"]
    question_options = list(range(len(questions)))
    selected_index = st.radio(
        "Choose a question to practice",
        options=question_options,
        format_func=lambda index: _question_picker_label(questions[index]),
    )

    selected_detail = questions[selected_index]
    selected_question = selected_detail["question"]
    st.markdown("**Practice question**")
    st.write(selected_question)
    st.caption(f"Rubric focus: {selected_detail['rubric_focus']}")

    answer = st.text_area(
        "Your answer",
        height=180,
        placeholder="Type a STAR-style answer here...",
        key="interview_answer_draft",
    )
    if answer.strip() and not st.session_state.interview_evaluation:
        advance_interview_progress("practice_started")

    if st.button("Evaluate my answer", type="primary"):
        answer_error = validate_interview_answer(answer)
        if answer_error:
            st.warning(answer_error)
        else:
            st.session_state.interview_evaluation = evaluate_interview_answer(
                selected_question,
                answer,
                selected_detail["rubric_focus"],
            )
            score = st.session_state.interview_evaluation["score"]
            advance_interview_progress(
                "ready_to_apply" if score >= 4 else "feedback_received",
                score,
            )

    if st.session_state.interview_evaluation:
        result = st.session_state.interview_evaluation
        st.subheader("Evaluation")
        st.metric("Practice score", f"{result['score']}/5")
        st.write(result["feedback"])
        with st.expander("Show a strong sample answer"):
            st.write(result["sample_answer"])


def _question_picker_label(question: dict) -> str:
    text = " ".join(str(question.get("question", "")).split())
    if len(text) > 92:
        text = text[:89].rsplit(" ", 1)[0].rstrip(",.;:") + "..."
    return f"{question.get('type', 'Question')}: {text}"


def render_report(analysis: dict) -> None:
    st.header("Final Navigation Brief")
    st.text_area("Copyable report", value=analysis["final_strategy_report"], height=360)

    st.warning(
        "CareerCompass recommendations are decision support, not a guarantee. Students should review outputs with a career advisor."
    )

    render_responsible_ai_audit(analysis)

    st.subheader("Rubric evidence snapshot")
    st.dataframe(pd.DataFrame(RUBRIC_EVIDENCE), use_container_width=True, hide_index=True)


def render_action_checklist(analysis: dict) -> None:
    st.header("Action Checklist")
    st.caption("Use this as a reminder system for missing resume evidence, certifications, and portfolio proof.")

    st.subheader("Resume evidence reminders")
    for item in analysis["resume_checklist"]:
        st.checkbox(item, value=False, key=f"resume-check-{item}")

    st.subheader("Suggested certifications and short courses")
    st.dataframe(
        pd.DataFrame(analysis["certification_recommendations"]),
        use_container_width=True,
        hide_index=True,
    )

    st.subheader("Portfolio proof ideas")
    for item in analysis["portfolio_ideas"]:
        st.checkbox(item, value=False, key=f"portfolio-{item}")


def render_rubric_evidence(analysis: dict) -> None:
    st.header("Rubric Evidence")
    st.caption("A compact proof map for the ISYS 573 Track B final submission.")

    st.subheader("Project requirement coverage")
    st.dataframe(pd.DataFrame(RUBRIC_EVIDENCE), use_container_width=True, hide_index=True)

    st.subheader("Agentic workflow proof")
    st.dataframe(
        pd.DataFrame(build_agent_coordination_rows(analysis)),
        use_container_width=True,
        hide_index=True,
    )

    st.subheader("RAG readiness")
    st.dataframe(pd.DataFrame(RAG_READINESS_ROWS), use_container_width=True, hide_index=True)

    st.subheader("Responsible AI controls")
    st.dataframe(pd.DataFrame(RESPONSIBLE_AI_CHECKS), use_container_width=True, hide_index=True)

    st.subheader("Submission checklist")
    checklist_items = [
        "Root tests pass with plain pytest discovery.",
        "Docker build and run commands are documented and ready for smoke testing.",
        "Screenshots are listed in docs/demo_flow.md for the final media package.",
        "Measured evaluation results are recorded in docs/evaluation_results.md.",
        "Final report and video can point to this table as rubric evidence.",
    ]
    for item in checklist_items:
        st.checkbox(item, value=True, key=f"rubric-check-{item}")


def render_responsible_ai_audit(analysis: dict) -> None:
    st.subheader("Responsible AI / advisor review audit")

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Resume retention", "Session only")
    c2.metric("Retrieved evidence", len(analysis.get("retrieved_job_postings", [])))
    c3.metric("Advisor review", "Required")
    c4.metric("Market confidence", _format_confidence(analysis, "market_data"))

    st.dataframe(pd.DataFrame(RESPONSIBLE_AI_CHECKS), use_container_width=True, hide_index=True)


def build_agent_coordination_rows(analysis: dict) -> list[dict]:
    handoffs = analysis.get("agent_handoffs", [])
    if handoffs:
        rows = []
        for index, handoff in enumerate(handoffs, start=1):
            rows.append(
                {
                    "Step": index,
                    "Coordination": f"{handoff['from_agent']} -> {handoff['to_agent']}",
                    "Reason": handoff["reason"],
                    "Inputs": ", ".join(handoff["required_inputs"]),
                    "Outputs": ", ".join(handoff["expected_outputs"]),
                    "Status": handoff["status"],
                }
            )
        return rows

    trace = analysis.get("agent_trace", [])
    return [
        {
            "Step": index,
            "Coordination": agent_name,
            "Reason": "Completed workflow step.",
            "Inputs": "Shared AgentState",
            "Outputs": "Updated AgentState",
            "Status": "completed",
        }
        for index, agent_name in enumerate(trace, start=1)
    ]


def _format_confidence(analysis: dict, key: str) -> str:
    value = analysis.get("confidence_scores", {}).get(key)
    if value is None:
        return "n/a"
    return f"{round(value * 100)}%"


def render_console_hero() -> None:
    analysis = st.session_state.get("analysis") or {}
    market_skills = analysis.get("market_skills", []) if analysis else []
    top_market_skill = market_skills[0].get("Skill", "Market signal") if market_skills else "Role signal"
    market_demand = market_skills[0].get("Demand Signal", "Pending") if market_skills else "Pending"
    gap_count = analysis.get("gap_counts", {}).get("high", 0) if analysis else 0
    next_action_count = len(analysis.get("next_actions", [])) if analysis else 0

    if analysis:
        metric_cards = [
            ("Target Role Fit", f"{analysis.get('match_percentage', '--')}%", analysis.get("role_label", "Role fit mapped"), ""),
            ("Skill Coverage", f"{analysis.get('keyword_coverage', '--')}%", "Resume evidence scan", ""),
            ("Market Demand", market_demand, top_market_skill, ""),
            ("Action Items", str(next_action_count), f"{gap_count} priority alerts", " is-alert"),
        ]
    else:
        metric_cards = [
            ("Target Role Fit", "Ready", "Load a profile", ""),
            ("Skill Coverage", "Scan", "Resume evidence", ""),
            ("Market Demand", "Pending", "Role signal", ""),
            ("Action Items", "0", "Generated after scan", " is-alert"),
        ]

    metrics_html = "".join(
        f"""
        <div class="cc-hero-metric{css_class}">
            <span>{escape(str(label))}</span>
            <strong>{escape(str(value))}</strong>
            <em>{escape(str(note))}</em>
        </div>
        """
        for label, value, note, css_class in metric_cards
    )

    st.markdown(
        f"""
        <div class="cc-hero">
            <div class="cc-hero-grid">
                <div>
                    <div class="cc-hero-kicker">Career Navigation Console</div>
                    <h1>CareerCompass</h1>
                    <p>Agentic AI career strategy MVP. Navigate roles, close skill gaps, rebuild resume evidence, and prepare interview routes from one command center.</p>
                </div>
                <div class="cc-compass-mark" aria-hidden="true"></div>
            </div>
            <div class="cc-chip-row">
                <span class="cc-chip">Profile scan ready</span>
                <span class="cc-chip">Market signal mapped</span>
                <span class="cc-chip">Resume targeting active</span>
                <span class="cc-chip">Interview route prepared</span>
            </div>
            <div class="cc-hero-metrics">{metrics_html}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def main() -> None:
    apply_visual_style()
    initialize_state()
    render_sidebar()
    apply_demo_query_params()

    render_console_hero()

    inputs = render_inputs()

    if st.button("Analyze resume and job fit", type="primary"):
        validation_error = validate_resume_for_analysis(inputs["resume_text"])
        if validation_error:
            st.session_state.analysis_validation_error = validation_error
            st.warning(validation_error)
        else:
            st.session_state.analysis_validation_error = ""
            started_at = time.perf_counter()
            render_progress()
            st.session_state.analysis = run_career_analysis(inputs)
            st.session_state.latency_seconds = time.perf_counter() - started_at
            reset_interview_progress()

    if st.session_state.analysis_validation_error and not st.session_state.analysis:
        st.error(st.session_state.analysis_validation_error)

    if st.session_state.analysis:
        views = [
            "Dashboard",
            "Roadmap",
            "Resume",
            "Interview",
            "Final Report",
            "Action Checklist",
            "Rubric Evidence",
        ]
        active_view = st.radio(
            "CareerCompass navigation",
            options=views,
            horizontal=True,
            key="active_view",
        )
        if active_view == "Dashboard":
            render_dashboard(st.session_state.analysis)
        elif active_view == "Roadmap":
            render_roadmap(st.session_state.analysis)
        elif active_view == "Resume":
            render_resume(st.session_state.analysis)
        elif active_view == "Interview":
            render_interview(st.session_state.analysis)
        elif active_view == "Final Report":
            render_report(st.session_state.analysis)
        elif active_view == "Action Checklist":
            render_action_checklist(st.session_state.analysis)
        elif active_view == "Rubric Evidence":
            render_rubric_evidence(st.session_state.analysis)


JOB_KEYWORD_ALIASES = {
    "SQL": ["sql", "query", "queries", "database"],
    "Python": ["python", "pandas", "notebook"],
    "Excel": ["excel", "spreadsheet", "pivot"],
    "Tableau": ["tableau"],
    "Power BI": ["power bi", "powerbi"],
    "Dashboarding": ["dashboard", "dashboards", "reporting", "visualization"],
    "Data analysis": ["data analysis", "analytics", "insights", "trend"],
    "A/B testing": ["a/b testing", "ab testing", "a/b test", "ab test", "experiment", "experimentation"],
    "Campaign performance": ["campaign performance", "campaign metrics", "marketing metrics"],
    "Content calendar": ["content calendar", "editorial calendar"],
    "Conversion funnel": ["conversion funnel", "conversion optimization", "funnel"],
    "Figma": ["figma"],
    "Go-to-market": ["go-to-market", "go to market", "gtm"],
    "Launch copy": ["launch copy", "launch messaging", "launch content"],
    "Positioning": ["positioning", "value proposition", "messaging"],
    "Requirements gathering": ["requirements", "user stories", "acceptance criteria"],
    "Stakeholder communication": ["stakeholder", "cross-functional", "presentation", "executive"],
    "Project management": ["project management", "timeline", "milestone", "delivery"],
    "Agile / Scrum": ["agile", "scrum", "sprint", "backlog"],
    "Risk management": ["risk", "mitigation", "blocker", "dependency"],
    "Customer insight": ["customer insight", "customer research", "user research", "persona", "voice of customer"],
    "Product marketing": ["product marketing", "launch", "campaign", "positioning"],
    "Process improvement": ["process improvement", "workflow", "automation", "operational"],
}


GENERIC_JOB_KEYWORD_WORDS = {
    "analysis",
    "applicant",
    "associate",
    "business",
    "candidate",
    "company",
    "customer",
    "experience",
    "marketing",
    "product",
    "requirements",
    "responsibilities",
    "role",
    "skills",
    "team",
    "work",
}


RESUME_SECTION_ALIASES = {
    "SUMMARY": "SUMMARY",
    "PROFILE": "SUMMARY",
    "PROFESSIONAL SUMMARY": "SUMMARY",
    "OBJECTIVE": "SUMMARY",
    "SKILLS": "SKILLS",
    "TECHNICAL SKILLS": "SKILLS",
    "CORE COMPETENCIES": "SKILLS",
    "EDUCATION": "EDUCATION",
    "ACADEMIC BACKGROUND": "EDUCATION",
    "EXPERIENCE": "EXPERIENCE",
    "WORK EXPERIENCE": "EXPERIENCE",
    "PROFESSIONAL EXPERIENCE": "EXPERIENCE",
    "EMPLOYMENT": "EXPERIENCE",
    "MILITARY EXPERIENCE": "EXPERIENCE",
    "LEADERSHIP EXPERIENCE": "EXPERIENCE",
    "PROJECTS": "PROJECTS",
    "PROJECT EXPERIENCE": "PROJECTS",
    "SELECTED PROJECTS": "PROJECTS",
    "CERTIFICATIONS": "CERTIFICATIONS",
    "CERTIFICATION": "CERTIFICATIONS",
    "CLEARANCE": "CERTIFICATIONS",
    "AWARDS": "ADDITIONAL",
    "VOLUNTEER": "ADDITIONAL",
    "LANGUAGES": "ADDITIONAL",
}


def resume_format_template_key(format_name: str) -> str:
    mapping = {
        "Use my existing resume structure": "ATS chronological",
        "Project-forward resume": "Project-forward",
        "Experience-forward resume": "ATS chronological",
        "Skills matrix resume": "Skills matrix",
        "Custom student-supplied template": "Use my own template",
    }
    return mapping.get(format_name, "ATS chronological")


def resume_format_metadata(format_name: str, analysis: dict, template_name: str) -> dict:
    if format_name == "Use my existing resume structure":
        return {
            "best_for": "Keeping the saved resume familiar while retargeting summary, skills, and bullets to the pasted job post.",
            "use_when": "Your current resume is already organized and you mainly need job-specific language.",
            "avoid_when": "Your experience is buried, missing sections, or hard to scan.",
            "sections": ["Header", "Summary", "Education", "Experience or Projects", "Skills", "Certifications"],
        }
    if format_name == "Experience-forward resume":
        return {
            "best_for": "Applications where the posting emphasizes ownership, impact, collaboration, or repeated work experience.",
            "use_when": "Your strongest evidence comes from jobs, military service, internships, or repeated responsibilities.",
            "avoid_when": "Your best proof is mostly class, portfolio, or capstone work.",
            "sections": ["Header", "Targeted Summary", "Core Skills", "Experience", "Projects", "Education"],
        }
    if template_name in analysis["resume_templates"]:
        template = dict(analysis["resume_templates"][template_name])
        if template_name == "Project-forward":
            template.setdefault("use_when", "Your strongest proof is project, portfolio, capstone, or class work.")
            template.setdefault("avoid_when", "A job post asks for deep professional experience first.")
        elif template_name == "Skills matrix":
            template.setdefault("use_when", "You are changing fields and need transferable evidence to be obvious.")
            template.setdefault("avoid_when", "The employer expects a conventional chronological resume.")
        else:
            template.setdefault("use_when", "You want a conventional ATS-friendly structure.")
            template.setdefault("avoid_when", "Your projects are stronger than your job history.")
        return template
    return {
        "best_for": "A targeted resume structure for the pasted job post.",
        "use_when": "You want a balanced layout.",
        "avoid_when": "You need a specialized template.",
        "sections": ["Header", "Summary", "Skills", "Projects", "Education"],
    }


def derive_job_post_keywords(job_post: str, analysis: dict, limit: int = 12) -> list[dict]:
    text = normalize_text(job_post)
    if not text:
        return []

    rows = []
    existing_keywords = [item["Keyword"] for item in analysis.get("keyword_targets", [])]
    candidate_labels = [*existing_keywords, *JOB_KEYWORD_ALIASES.keys()]
    seen = set()

    for label in candidate_labels:
        normalized_label = label.lower()
        if normalized_label in seen:
            continue
        aliases = JOB_KEYWORD_ALIASES.get(label, [label])
        if any(alias.lower() in text for alias in aliases):
            rows.append(
                {
                    "Keyword": label,
                    "Source": "Job post",
                    "How to use it": job_keyword_usage(label),
                }
            )
            seen.add(normalized_label)

    for phrase in extract_repeated_job_phrases(job_post):
        normalized_phrase = phrase.lower()
        if normalized_phrase not in seen:
            rows.append(
                {
                    "Keyword": phrase,
                    "Source": "Repeated phrase",
                    "How to use it": "Mirror this phrase only where your saved resume has matching evidence.",
                }
            )
            seen.add(normalized_phrase)
        if len(rows) >= limit:
            break

    return rows[:limit]


def build_skill_evidence_map(
    analysis: dict,
    resume_text: str,
    coursework: list[str],
    job_post: str = "",
) -> list[dict]:
    job_keywords = derive_job_post_keywords(job_post, analysis)
    job_keyword_names = [item["Keyword"] for item in job_keywords]
    market_skill_demand = {
        item["Skill"]: item.get("Demand Signal", "Medium")
        for item in analysis.get("market_skills", [])
    }
    candidate_skills = merge_unique(
        job_keyword_names,
        [
            *[item["Keyword"] for item in analysis.get("keyword_targets", [])],
            *market_skill_demand.keys(),
        ],
    )[:14]

    rows = []
    for skill in candidate_skills:
        evidence = assess_skill_evidence(skill, resume_text, coursework)
        demand = "Job post" if skill in job_keyword_names else market_skill_demand.get(skill, "Profile target")
        rows.append(
            {
                "Skill": skill,
                "Job Demand": demand,
                "Resume Evidence": evidence_status_label(evidence["status"]),
                "Found In": evidence["evidence_source"],
                "Evidence Excerpt": evidence["evidence_excerpt"],
                "Gap Action": skill_gap_action(skill, evidence["status"], demand),
            }
        )
    return rows


def evidence_status_label(status: str) -> str:
    return {
        "Strong Evidence": "Strong evidence",
        "Mentioned": "Mentioned only",
        "Missing": "Needs proof",
    }.get(status, status)


def skill_gap_action(skill: str, status: str, demand: str) -> str:
    if status == "Strong Evidence":
        return "Keep it visible and quantify the outcome if possible."
    if status == "Mentioned":
        return f"Add a bullet showing how you used {skill} in a project, job, or coursework deliverable."
    if demand == "Job post":
        return f"Add truthful {skill} evidence or avoid over-claiming it for this application."
    return f"Build or document one small proof point for {skill}."


def job_keyword_usage(keyword: str) -> str:
    lower_keyword = keyword.lower()
    if any(term in lower_keyword for term in ["sql", "python", "excel", "tableau", "power bi", "dashboard"]):
        return "Name the tool near a project, analysis task, or measurable deliverable."
    if any(term in lower_keyword for term in ["stakeholder", "requirements", "project", "agile", "risk"]):
        return "Use it in an impact bullet that shows coordination, decisions, or delivery."
    if any(term in lower_keyword for term in ["marketing", "customer", "testing", "process"]):
        return "Connect it to customer insight, experiments, workflow improvement, or launch outcomes."
    return "Place it in the summary, skills, or a bullet only when supported by your resume evidence."


def extract_repeated_job_phrases(job_post: str) -> list[str]:
    normalized = normalize_text(job_post)
    words = re.findall(r"[A-Za-z][A-Za-z+#/.-]{2,}", normalized)
    stop_words = {
        "and",
        "the",
        "for",
        "with",
        "you",
        "our",
        "will",
        "this",
        "that",
        "from",
        "are",
        "have",
        "your",
        "work",
        "team",
        "role",
        "skills",
        "experience",
        "ability",
        *GENERIC_JOB_KEYWORD_WORDS,
    }
    phrases = []
    for label, aliases in JOB_KEYWORD_ALIASES.items():
        if label in {"SQL", "Python", "Excel", "Tableau", "Power BI"}:
            continue
        if any(alias.lower() in normalized for alias in aliases):
            phrases.append(label)

    phrase_counts = {}
    for size in (3, 2):
        for index in range(0, max(len(words) - size + 1, 0)):
            chunk = words[index : index + size]
            if any(word in stop_words or len(word) < 4 for word in chunk):
                continue
            phrase = " ".join(chunk)
            phrase_counts[phrase] = phrase_counts.get(phrase, 0) + 1

    ranked_phrases = sorted(phrase_counts.items(), key=lambda item: (-item[1], item[0]))
    for phrase, count in ranked_phrases:
        if count >= 2 and phrase not in {item.lower() for item in phrases}:
            phrases.append(phrase)

    counts = {}
    for word in words:
        if word in stop_words or len(word) < 4:
            continue
        counts[word] = counts.get(word, 0) + 1

    ranked = sorted(counts.items(), key=lambda item: (-item[1], item[0]))
    for word, count in ranked:
        if count >= 3 and word not in {item.lower() for item in phrases}:
            phrases.append(word.title() if word in {"sql", "python", "tableau"} else word)

    return phrases[:6]


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.lower()).strip()


def infer_job_title(job_post: str, analysis: dict, target_job: dict[str, str] | None = None) -> str:
    if target_job and target_job.get("title"):
        return clean_job_title(target_job["title"])
    lines = [line.strip(" -|") for line in job_post.splitlines() if line.strip()]
    for line in lines[:6]:
        if 3 <= len(line) <= 80 and not line.endswith(".") and not line.lower().startswith(("http://", "https://")):
            return clean_job_title(line)
    return analysis.get("target_role") or analysis.get("role_label") or "Target job"


def clean_job_title(title: str) -> str:
    title = re.sub(r"\s+", " ", title).strip(" -|")
    parts = [part.strip() for part in re.split(r"\s+-\s+|\s+\|\s+", title) if part.strip()]
    if parts:
        return parts[0]
    return title


def parse_resume_sections(saved_resume: str) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {"HEADER": []}
    current = "HEADER"

    for raw_line in saved_resume.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        heading, remainder = split_resume_heading(line)
        if heading:
            current = heading
            sections.setdefault(current, [])
            if remainder:
                sections[current].append(remainder)
            continue

        sections.setdefault(current, []).append(line)

    return sections


def split_resume_heading(line: str) -> tuple[str | None, str]:
    normalized = re.sub(r"[^A-Za-z ]", "", line).strip().upper()
    if normalized in RESUME_SECTION_ALIASES:
        return RESUME_SECTION_ALIASES[normalized], ""

    if ":" in line:
        prefix, remainder = line.split(":", 1)
        normalized_prefix = re.sub(r"[^A-Za-z ]", "", prefix).strip().upper()
        if normalized_prefix in RESUME_SECTION_ALIASES:
            return RESUME_SECTION_ALIASES[normalized_prefix], remainder.strip()

    return None, ""


def extract_resume_identity(saved_resume: str) -> dict[str, str]:
    sections = parse_resume_sections(saved_resume)
    header_lines = sections.get("HEADER", [])
    fallback_lines = [line.strip() for line in saved_resume.splitlines() if line.strip()]
    candidate_lines = header_lines or fallback_lines[:6]

    email = find_first(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}", saved_resume)
    phone = find_first(r"(?:\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", saved_resume)
    links = re.findall(r"(?:https?://)?(?:www\.)?(?:linkedin\.com|github\.com|[\w.-]+\.(?:com|net|io|me))/\S+", saved_resume, flags=re.IGNORECASE)

    name = "NAME"
    for line in candidate_lines:
        if line == email or line == phone:
            continue
        if email and email.lower() in line.lower():
            continue
        if phone and phone in line:
            continue
        if not has_contact_signal(line) and len(line.split()) <= 5:
            name = line.strip(" -")
            break

    location = ""
    for line in candidate_lines:
        if line == name or has_contact_signal(line):
            continue
        if any(char.isdigit() for char in line) and "," in line:
            location = line
            break

    contact_parts = [part for part in [email, phone, *links[:2], location] if part]
    return {
        "name": name,
        "contact": " | ".join(dict.fromkeys(contact_parts)),
    }


def find_first(pattern: str, text: str) -> str:
    match = re.search(pattern, text, flags=re.IGNORECASE)
    return match.group(0) if match else ""


def has_contact_signal(line: str) -> bool:
    lower_line = line.lower()
    return bool(
        "@" in line
        or re.search(r"\d{3}[-.\s)]?\d{3}[-.\s]?\d{4}", line)
        or "linkedin" in lower_line
        or "github" in lower_line
        or "portfolio" in lower_line
        or "http" in lower_line
    )


def extract_resume_education(saved_resume: str) -> list[str]:
    sections = parse_resume_sections(saved_resume)
    education = clean_resume_lines(sections.get("EDUCATION", []))
    if education:
        return education

    degree_signals = ["degree", "bachelor", "master", "associate", "college", "university", "school", "sfsu"]
    return [
        line.strip()
        for line in saved_resume.splitlines()
        if any(signal in line.lower() for signal in degree_signals)
    ][:8]


def extract_resume_skills(saved_resume: str) -> list[str]:
    sections = parse_resume_sections(saved_resume)
    skills = clean_resume_lines(sections.get("SKILLS", []))
    return split_inline_items(skills)[:16]


def extract_resume_bullets(saved_resume: str, limit: int = 6) -> list[str]:
    sections = parse_resume_sections(saved_resume)
    source_lines = []
    for section in ["EXPERIENCE", "PROJECTS", "SUMMARY", "ADDITIONAL"]:
        source_lines.extend(sections.get(section, []))

    bullets = []
    for line in clean_resume_lines(source_lines):
        if has_contact_signal(line):
            continue
        if looks_like_education_line(line):
            continue
        if split_resume_heading(line)[0]:
            continue
        bullets.extend(split_resume_line_into_bullets(line))

    return bullets[:limit]


def clean_resume_lines(lines: list[str]) -> list[str]:
    cleaned = []
    for line in lines:
        stripped = line.strip().lstrip("-*• ").strip()
        if stripped and stripped not in cleaned:
            cleaned.append(stripped)
    return cleaned


def split_inline_items(lines: list[str]) -> list[str]:
    items = []
    for line in lines:
        chunks = re.split(r",|;|\|", line)
        for chunk in chunks:
            item = clean_skill_label(chunk)
            if item and item not in items:
                items.append(item)
    return items or lines


def split_resume_line_into_bullets(line: str) -> list[str]:
    stripped = line.strip().lstrip("-*• ").strip()
    if len(stripped) > 220:
        parts = [part.strip() for part in re.split(r"(?<=[.!?])\s+", stripped) if len(part.strip()) > 35]
        return parts or [stripped[:217].rstrip(",.;:") + "..."]
    return [stripped]


def looks_like_education_line(line: str) -> bool:
    lower_line = line.lower()
    return any(
        signal in lower_line
        for signal in ["bachelor", "master", "associate", "college", "university", "degree", "sfsu", "gpa"]
    )


def keyword_names(analysis: dict, job_post: str) -> list[str]:
    names = [item["Keyword"] for item in derive_job_post_keywords(job_post, analysis)]
    for item in analysis.get("keyword_targets", []):
        if item["Keyword"] not in names:
            names.append(item["Keyword"])
    return names[:10]


def build_tailoring_change_summary(
    saved_resume: str,
    analysis: dict,
    job_post: str,
    selected_suggestions: list[dict] | None = None,
    coursework: list[str] | None = None,
    safe_edits_only: bool = False,
) -> list[dict]:
    coursework = coursework or []
    rows = []
    resume_lower = normalize_text(saved_resume)
    suggestions = selected_suggestions or analysis.get("resume_recommendations", [])[:4]

    for suggestion in suggestions:
        keywords = suggestion.get("keywords_added", [])
        evidence_statuses = [
            assess_skill_evidence(keyword, saved_resume, coursework)["status"]
            for keyword in keywords
        ]
        before = suggestion.get("before", "")
        before_in_resume = bool(before and normalize_text(before) in resume_lower)
        safe = any(
            evidence_status_allows_body(status, safe_edits_only)
            for status in evidence_statuses
        ) or (before_in_resume and (not safe_edits_only or not evidence_statuses))
        rows.append(
            {
                "Original": before or "Resume evidence not found",
                "Suggested": suggestion.get("after", "") if safe else f"Needs source evidence before adding: {suggestion.get('after', '')}",
                "Reason": "Selected resume improvement" if safe else "Suggestion is useful but not clearly supported by the supplied resume.",
                "Evidence Status": ", ".join(evidence_statuses) if evidence_statuses else "Not checked",
                "Safe To Add": "Yes" if safe else "Needs evidence",
            }
        )

    for keyword in keyword_names(analysis, job_post):
        evidence = assess_skill_evidence(keyword, saved_resume, coursework)
        if evidence["status"] == "Missing":
            rows.append(
                {
                    "Original": "No resume evidence found",
                    "Suggested": f"Do not claim {keyword} yet; add a project, coursework, or work example first.",
                    "Reason": "Keyword appears in the target job but is not supported by current resume evidence.",
                    "Evidence Status": "Missing",
                    "Safe To Add": "Needs evidence",
                }
            )
    return rows[:12]


def summarize_tailoring_result(change_rows: list[dict], job_keywords: list[dict]) -> dict[str, int]:
    return {
        "keywords": len(job_keywords),
        "revised_bullets": sum(1 for row in change_rows if row.get("Safe To Add") == "Yes"),
        "missing_skills": sum(1 for row in change_rows if row.get("Safe To Add") != "Yes"),
    }


def evidence_status_allows_body(status: str, safe_edits_only: bool = False) -> bool:
    if status == "Strong Evidence":
        return True
    if status == "Mentioned":
        return not safe_edits_only
    return False


def suggestion_is_supported(
    suggestion: dict,
    saved_resume: str,
    coursework: list[str],
    safe_edits_only: bool = False,
) -> bool:
    before = normalize_text(suggestion.get("before", ""))
    resume_lower = normalize_text(saved_resume)
    if before and before in resume_lower and not safe_edits_only:
        return True
    return any(
        evidence_status_allows_body(
            assess_skill_evidence(keyword, saved_resume, coursework)["status"],
            safe_edits_only,
        )
        for keyword in suggestion.get("keywords_added", [])
    )


def build_tailored_resume_draft(
    analysis: dict,
    saved_resume: str,
    job_post: str,
    format_name: str,
    custom_template: str = "",
    selected_suggestions: list[dict] | None = None,
    coursework: list[str] | None = None,
    safe_edits_only: bool = False,
    regeneration_mode: str = "Standard",
    target_job: dict[str, str] | None = None,
) -> str:
    coursework = coursework or []
    role = analysis.get("role_label", analysis.get("target_role", "Target Role"))
    job_title = infer_job_title(job_post, analysis, target_job)
    company = (target_job or {}).get("company", "").strip()
    keywords = keyword_names(analysis, job_post)
    evidence_by_keyword = {
        keyword: assess_skill_evidence(keyword, saved_resume, coursework)
        for keyword in keywords
    }
    supported_keywords = [
        keyword
        for keyword, evidence in evidence_by_keyword.items()
        if evidence_status_allows_body(evidence["status"], safe_edits_only)
    ]
    missing_keywords = [
        keyword
        for keyword, evidence in evidence_by_keyword.items()
        if not evidence_status_allows_body(evidence["status"], safe_edits_only)
    ]
    keyword_line = ", ".join((supported_keywords or keywords)[:8])
    identity = extract_resume_identity(saved_resume)
    education_lines = extract_resume_education(saved_resume)
    saved_skills = extract_resume_skills(saved_resume)
    source_bullets = extract_resume_bullets(saved_resume)
    candidate_suggestions = selected_suggestions or analysis.get("resume_recommendations", [])[:4]
    safe_suggestions = [
        item
        for item in candidate_suggestions
        if suggestion_is_supported(item, saved_resume, coursework, safe_edits_only)
    ]
    suggestion_limit = 2 if regeneration_mode == "More concise" else 4
    suggestion_bullets = [item["after"] for item in safe_suggestions[:suggestion_limit]]

    targeted_bullets = tailor_bullets(source_bullets, suggestion_bullets)
    candidate_label = f"{role} candidate" if role.lower() in job_title.lower() else "candidate"
    target_phrase = f"for {job_title}" if job_title else f"for {role}"
    company_phrase = f" at {company}" if company else ""
    if regeneration_mode == "More concise":
        summary = f"Early-career {candidate_label} prepared {target_phrase}{company_phrase}, with supported evidence in {keyword_line or 'transferable experience'}."
    else:
        summary = (
            f"Early-career {candidate_label} prepared {target_phrase}{company_phrase}. "
            f"Resume evidence supports {keyword_line or 'relevant coursework and transferable experience'}, "
            "with unsupported claims separated for review."
        )
    header = identity["name"]
    if identity["contact"]:
        header += f"\n{identity['contact']}"
    education = "\n".join(education_lines) if education_lines else "Add education from saved resume."
    skills_with_resume = merge_unique(saved_skills, supported_keywords)[:14]
    skills = ", ".join(skills_with_resume)
    improvements = "\n".join(f"- {bullet}" for bullet in suggestion_bullets[:4]) or "- No unsupported rewrite was inserted automatically."
    needs_evidence = "\n".join(
        f"- {keyword}: add truthful project, coursework, or work evidence before claiming this in the resume body."
        for keyword in missing_keywords[:8]
    ) or "- No unsupported job-post keywords were flagged."
    source_evidence_limit = 3 if regeneration_mode == "More concise" else 4

    section_map = {
        "HEADER": header,
        "SUMMARY": f"SUMMARY\n{summary}",
        "TARGETED SKILLS": f"TARGETED SKILLS\n{skills}",
        "EXPERIENCE": "EXPERIENCE\n" + "\n".join(f"- {bullet}" for bullet in targeted_bullets),
        "PROJECT EXPERIENCE": "PROJECT EXPERIENCE\n" + "\n".join(f"- {bullet}" for bullet in targeted_bullets),
        "EDUCATION": f"EDUCATION\n{education}",
        "SOURCE RESUME EVIDENCE": "SOURCE RESUME EVIDENCE\n" + "\n".join(f"- {bullet}" for bullet in source_bullets[:source_evidence_limit]),
        "CAREERCOMPASS IMPROVEMENTS": f"CAREERCOMPASS IMPROVEMENTS\n{improvements}",
        "DO NOT ADD UNTIL YOU HAVE EVIDENCE": f"DO NOT ADD UNTIL YOU HAVE EVIDENCE\n{needs_evidence}",
        "JOB POST TARGETS": "JOB POST TARGETS\n" + "\n".join(f"- {keyword}" for keyword in keywords[:8]),
    }

    if format_name == "Custom student-supplied template" and custom_template.strip():
        custom_sections = [line.strip().upper() for line in custom_template.splitlines() if line.strip()]
        rendered = [section_map["HEADER"]]
        for section in custom_sections:
            rendered.append(section_map.get(section, f"{section}\nAdd content here using saved resume evidence and job-post keywords."))
        rendered.append(section_map["CAREERCOMPASS IMPROVEMENTS"])
        rendered.append(section_map["DO NOT ADD UNTIL YOU HAVE EVIDENCE"])
        rendered.append(section_map["JOB POST TARGETS"])
        return "\n\n".join(rendered).strip() + "\n"

    if regeneration_mode == "ATS-focused":
        order = ["HEADER", "SUMMARY", "TARGETED SKILLS", "EXPERIENCE", "EDUCATION", "CAREERCOMPASS IMPROVEMENTS", "DO NOT ADD UNTIL YOU HAVE EVIDENCE", "JOB POST TARGETS"]
    elif format_name == "Project-forward resume":
        order = ["HEADER", "SUMMARY", "TARGETED SKILLS", "PROJECT EXPERIENCE", "EDUCATION", "CAREERCOMPASS IMPROVEMENTS", "DO NOT ADD UNTIL YOU HAVE EVIDENCE", "JOB POST TARGETS"]
    elif format_name == "Experience-forward resume":
        order = ["HEADER", "SUMMARY", "TARGETED SKILLS", "EXPERIENCE", "EDUCATION", "CAREERCOMPASS IMPROVEMENTS", "DO NOT ADD UNTIL YOU HAVE EVIDENCE"]
    elif format_name == "Skills matrix resume":
        order = ["HEADER", "SUMMARY", "TARGETED SKILLS", "SOURCE RESUME EVIDENCE", "PROJECT EXPERIENCE", "EDUCATION", "CAREERCOMPASS IMPROVEMENTS", "DO NOT ADD UNTIL YOU HAVE EVIDENCE"]
    else:
        order = ["HEADER", "SUMMARY", "EXPERIENCE", "EDUCATION", "TARGETED SKILLS", "CAREERCOMPASS IMPROVEMENTS", "DO NOT ADD UNTIL YOU HAVE EVIDENCE", "JOB POST TARGETS"]

    return "\n\n".join(section_map[section] for section in order).strip() + "\n"


def tailor_bullets(source_bullets: list[str], suggestion_bullets: list[str]) -> list[str]:
    combined = [*source_bullets[:4], *suggestion_bullets[:3]]
    if not combined:
        combined = ["Add source resume bullets before exporting this draft."]
    return combined[:6]


def merge_unique(primary: list[str], secondary: list[str]) -> list[str]:
    merged = []
    for item in [*primary, *secondary]:
        cleaned = clean_skill_label(item)
        if cleaned and cleaned.lower() not in {existing.lower() for existing in merged}:
            merged.append(cleaned)
    return merged


def clean_skill_label(item: str) -> str:
    return re.sub(r"\s+", " ", item).strip().strip(" .;:,")


def build_resume_draft(
    analysis: dict,
    template_name: str = "ATS chronological",
    custom_template: str = "",
) -> str:
    role = analysis.get("role_label", analysis.get("target_role", "Target Role"))
    keywords = ", ".join(item["Keyword"] for item in analysis["keyword_targets"][:6])
    bullets = "\n".join(
        f"- {item['after']}" for item in analysis["resume_recommendations"]
    )
    certifications = "\n".join(
        f"- {item['Recommendation']} ({item['Priority']})"
        for item in analysis["certification_recommendations"][:4]
    )

    if template_name == "Use my own template" and custom_template.strip():
        return (
            "NAME\n"
            "Email | Phone | LinkedIn | Portfolio\n\n"
            "CUSTOM TEMPLATE STRUCTURE\n"
            f"{custom_template.strip()}\n\n"
            f"TARGET ROLE\n{role}\n\n"
            f"KEYWORD TARGETS\n{keywords}\n\n"
            f"SELECTED EXPERIENCE BULLETS\n{bullets}\n\n"
            f"CERTIFICATIONS TO ADD OR PURSUE\n{certifications}\n"
        )

    if template_name == "Project-forward":
        return (
            "NAME\n"
            "Email | Phone | LinkedIn | Portfolio\n\n"
            f"TARGET ROLE\n{role}\n\n"
            "SELECTED PROJECTS\n"
            "CareerCompass Portfolio Project\n"
            f"{bullets}\n\n"
            f"SKILLS SNAPSHOT\n{keywords}\n\n"
            "EDUCATION\n"
            "B.S. Management Information Systems\n\n"
            f"CERTIFICATIONS\n{certifications}\n"
        )

    if template_name == "Skills matrix":
        return (
            "NAME\n"
            "Email | Phone | LinkedIn | Portfolio\n\n"
            "SUMMARY\n"
            f"Early-career candidate targeting {role} roles with evidence across analysis, tools, projects, and stakeholder communication.\n\n"
            "SKILLS MATRIX\n"
            f"Target Keywords: {keywords}\n"
            "Business Skills: stakeholder communication, requirements, prioritization, recommendations\n"
            "Tools: SQL, dashboards, Excel, project tracking tools\n\n"
            "EVIDENCE BULLETS\n"
            f"{bullets}\n\n"
            "EDUCATION AND CERTIFICATIONS\n"
            "B.S. Management Information Systems\n"
            f"{certifications}\n"
        )

    return (
        "NAME\n"
        "Email | Phone | LinkedIn | Portfolio\n\n"
        "SUMMARY\n"
        f"Early-career {role} candidate with experience in analytics, project work, and stakeholder-focused problem solving. Target keywords include {keywords}.\n\n"
        "EDUCATION\n"
        "B.S. Management Information Systems\n\n"
        "PROJECT EXPERIENCE\n"
        f"{bullets}\n\n"
        "SKILLS\n"
        f"{keywords}\n\n"
        "CERTIFICATIONS\n"
        f"{certifications}\n"
    )


def apply_resume_suggestions(draft: str, suggestions: list[dict]) -> str:
    if not suggestions:
        return draft

    applied_lines = "\n".join(f"- {suggestion['after']}" for suggestion in suggestions)
    return (
        draft.rstrip()
        + "\n\nAPPLIED CAREERCOMPASS IMPROVEMENTS\n"
        + applied_lines
        + "\n"
    )


def build_docx_bytes(text: str) -> bytes:
    try:
        from docx import Document
    except ImportError:
        return text.encode("utf-8")

    document = Document()
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            document.add_paragraph("")
        elif line.isupper() and len(line) <= 48:
            document.add_heading(line.title(), level=2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        else:
            document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


if __name__ == "__main__":
    main()
